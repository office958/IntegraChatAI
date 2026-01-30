from fastapi import APIRouter, UploadFile, File, Form, Query
from fastapi.responses import JSONResponse
from typing import Optional
import io
from core.config import PDF_AVAILABLE, OCR_AVAILABLE
import PyPDF2
from PIL import Image
import pytesseract

# Importă python-docx pentru procesarea fișierelor DOCX
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# Importă OCR-ul nou cu PaddleOCR
try:
    from ocr_processor.processor import process_image, process_pdf, PADDLEOCR_AVAILABLE, OPENCV_AVAILABLE
    from ocr_processor.postprocess import correct_ocr_text, identify_missing_fields
    PADDLEOCR_AVAILABLE_IMPORT = PADDLEOCR_AVAILABLE and OPENCV_AVAILABLE
except ImportError:
    PADDLEOCR_AVAILABLE_IMPORT = False
    process_pdf = None
    # Importă funcțiile de post-procesare dacă sunt disponibile
    try:
        from ocr_processor.postprocess import correct_ocr_text, identify_missing_fields
    except ImportError:
        correct_ocr_text = None
        identify_missing_fields = None

router = APIRouter(tags=["files"])

# Importă PDF2IMAGE_AVAILABLE din config (verificat la start)
from core.config import PDF2IMAGE_AVAILABLE

# Importă convert_from_bytes dacă este disponibil
if PDF2IMAGE_AVAILABLE:
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        PDF2IMAGE_AVAILABLE = False

@router.post("/extract-pdf")
async def extract_pdf(
    pdf: UploadFile = File(...), 
    max_pages: int = Query(5, ge=1, le=10, description="Numărul maxim de pagini de procesat (1-10, default: 5)")
):
    """
    Extrage textul dintr-un fișier PDF.
    Dacă PDF-ul este scanat (fără text extractibil), folosește OCR ca fallback.
    
    Args:
        pdf: Fișierul PDF
        max_pages: Numărul maxim de pagini de procesat (default: 10, pentru a evita timeout-uri)
    """
    if not PDF_AVAILABLE:
        return JSONResponse(
            status_code=500,
            content={"error": "PyPDF2 nu este instalat. Rulează: pip install PyPDF2"}
        )
    
    if pdf.content_type != "application/pdf":
        return JSONResponse(
            status_code=400,
            content={"error": "Fișierul trebuie să fie PDF"}
        )
    
    try:
        # Citește conținutul PDF
        pdf_content = await pdf.read()
        
        if not pdf_content:
            return JSONResponse(
                status_code=400,
                content={"error": "Fișierul PDF este gol sau nu a putut fi citit."}
            )
        
        # Încearcă mai întâi extragerea textului direct cu PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        text = ""
        extracted_pages = 0
        
        # Limitează numărul de pagini pentru a evita timeout-uri
        total_pages = len(pdf_reader.pages)
        pages_to_process = min(total_pages, max_pages)
        
        for page_num, page in enumerate(pdf_reader.pages[:pages_to_process]):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += f"\n--- Pagina {page_num + 1} ---\n"
                    text += page_text
                    extracted_pages += 1
            except Exception as e:
                print(f"⚠️ Eroare la extragerea paginii {page_num + 1}: {e}")
                continue
        
        # Adaugă notă dacă au fost omise pagini
        if total_pages > max_pages and extracted_pages > 0:
            text += f"\n\n[Notă: Doar primele {max_pages} pagini au fost procesate din {total_pages} totale]"
        
        # Dacă nu s-a extras text sau s-a extras foarte puțin, încercă OCR (dacă este disponibil)
        if not text.strip() or (extracted_pages == 0 and len(pdf_reader.pages) > 0):
            print(f"📄 PDF pare să fie scanat (fără text extractibil). Încerc OCR...")
            
            if not PDF2IMAGE_AVAILABLE:
                return JSONResponse(
                    status_code=400,
                    content={
                        "error": "Nu s-a putut extrage text din PDF. PDF-ul pare să fie scanat. Pentru a procesa PDF-uri scanate, instalează: pip install pdf2image și poppler (Windows: https://github.com/oschwartz10612/poppler-windows/releases, Linux: sudo apt-get install poppler-utils, macOS: brew install poppler)"
                    }
                )
            
            # Preferă PaddleOCR dacă este disponibil, altfel folosește Tesseract
            if PADDLEOCR_AVAILABLE_IMPORT and process_pdf:
                print(f"📄 Folosesc PaddleOCR pentru extragere text din PDF scanat (max {max_pages} pagini)...")
                try:
                    # Rulează conversia PDF->imagini în thread pool pentru a nu bloca event loop-ul
                    import asyncio
                    loop = asyncio.get_event_loop()
                    # Convertește doar primele max_pages pagini pentru viteză maximă
                    # Folosim DPI mai mic (150) pentru viteză mai mare
                    images = await loop.run_in_executor(
                        None,
                        lambda: convert_from_bytes(pdf_content, dpi=150)
                    )
                    
                    if not images:
                        print(f"⚠️ Nu s-au putut converti paginile PDF, încerc cu Tesseract...")
                    else:
                        # Limitează la primele max_pages pagini
                        total_pages = len(images)
                        images_to_process = images[:max_pages]
                        
                        ocr_text = ""
                        processor = None
                        
                        # Folosește singleton pentru a evita reinițializarea modelelor
                        try:
                            from ocr_processor.singleton import get_ocr_processor
                            processor = get_ocr_processor(lang='ro')
                            if not processor:
                                raise Exception("OCRProcessor nu este disponibil")
                        except Exception as e:
                            print(f"⚠️ Eroare la obținerea OCRProcessor: {e}")
                            raise
                        
                        # Procesează fiecare pagină asincron (limitează la max_pages pentru viteză)
                        pages_processed = 0
                        for page_num, pil_image in enumerate(images_to_process):
                            if pages_processed >= max_pages:
                                break
                            
                            # Rulează procesarea în thread pool pentru a nu bloca
                            try:
                                page_text, _ = await loop.run_in_executor(
                                    None,
                                    lambda img=pil_image: processor.process_pil_image(img, return_boxes=False)
                                )
                                
                                if page_text and page_text.strip():
                                    ocr_text += f"\n--- Pagina {page_num + 1} ---\n"
                                    ocr_text += page_text
                                    pages_processed += 1
                                    print(f"✅ Text extras din pagina {page_num + 1} cu PaddleOCR: {len(page_text)} caractere")
                            except Exception as page_error:
                                import traceback
                                print(f"⚠️ Eroare la extragerea paginii {page_num + 1} cu PaddleOCR: {traceback.format_exc()}")
                                continue
                        
                        # Adaugă notă dacă au fost omise pagini
                        if total_pages > max_pages:
                            ocr_text += f"\n\n[Notă: Doar primele {pages_processed} pagini au fost procesate din {total_pages} totale]"
                        
                        if ocr_text.strip():
                            # Limitează textul la 50000 caractere pentru a evita timeout-uri
                            final_text = ocr_text.strip()
                            if len(final_text) > 50000:
                                final_text = final_text[:50000] + "\n\n[... text trunchiat pentru a evita timeout-uri ...]"
                            
                            return JSONResponse(content={
                                "text": final_text,
                                "filename": pdf.filename,
                                "type": "pdf",
                                "method": "paddleocr",
                                "pages": pages_processed,
                                "total_pages": total_pages,
                                "truncated": len(ocr_text.strip()) > 50000
                            })
                        else:
                            print(f"⚠️ PaddleOCR nu a extras text, încerc cu Tesseract...")
                except Exception as paddle_error:
                    import traceback
                    print(f"⚠️ Eroare la PaddleOCR pentru PDF: {traceback.format_exc()}")
                    print(f"⚠️ Încerc cu Tesseract ca fallback...")
            
            # Fallback la Tesseract dacă PaddleOCR nu este disponibil sau a eșuat
            if not OCR_AVAILABLE:
                return JSONResponse(
                    status_code=400,
                    content={
                        "error": "Nu s-a putut extrage text din PDF. PDF-ul pare să fie scanat. Pentru a procesa PDF-uri scanate, instalează OCR: pip install pytesseract pillow pdf2image și Tesseract OCR."
                    }
                )
            
            # Verifică dacă Tesseract funcționează
            try:
                pytesseract.get_tesseract_version()
            except Exception as tess_error:
                return JSONResponse(
                    status_code=500,
                    content={
                        "error": f"Tesseract OCR nu este disponibil. Eroare: {str(tess_error)}. Instalează Tesseract OCR de la: https://github.com/UB-Mannheim/tesseract/wiki"
                    }
                )
            
            # Convertește PDF-ul în imagini și extrage text cu OCR
            try:
                # Rulează conversia PDF->imagini în thread pool pentru a nu bloca event loop-ul
                import asyncio
                loop = asyncio.get_event_loop()
                # Reduce DPI-ul la 150 pentru viteză mai mare
                images = await loop.run_in_executor(
                    None,
                    lambda: convert_from_bytes(pdf_content, dpi=150)
                )
                
                # Limitează la primele max_pages pagini
                total_pages = len(images)
                images_to_process = images[:max_pages]
                
                ocr_text = ""
                
                for page_num, img in enumerate(images_to_process):
                    try:
                        # Încearcă cu diferite configurații de limbi
                        page_ocr_text = None
                        lang_configs = ['ron+eng', 'eng', 'ron', None]
                        
                        # Rulează OCR-ul în thread pool pentru a nu bloca event loop-ul
                        loop = asyncio.get_event_loop()
                        for lang_config in lang_configs:
                            try:
                                if lang_config:
                                    page_ocr_text = await loop.run_in_executor(
                                        None,
                                        lambda img=img, lang=lang_config: pytesseract.image_to_string(img, lang=lang)
                                    )
                                else:
                                    page_ocr_text = await loop.run_in_executor(
                                        None,
                                        lambda img=img: pytesseract.image_to_string(img)
                                    )
                                
                                if page_ocr_text and page_ocr_text.strip():
                                    break
                            except Exception as e:
                                if "tesseract" in str(e).lower() or "not found" in str(e).lower():
                                    raise e
                                continue
                        
                        if page_ocr_text and page_ocr_text.strip():
                            ocr_text += f"\n--- Pagina {page_num + 1} (OCR) ---\n"
                            ocr_text += page_ocr_text.strip()
                    except Exception as e:
                        print(f"⚠️ Eroare la OCR pentru pagina {page_num + 1}: {e}")
                        continue
                
                # Adaugă notă dacă au fost omise pagini
                if total_pages > max_pages:
                    ocr_text += f"\n\n[Notă: Doar primele {max_pages} pagini au fost procesate din {total_pages} totale]"
                
                if ocr_text.strip():
                    # Limitează textul la 50000 caractere pentru a evita timeout-uri
                    final_ocr_text = ocr_text.strip()
                    if len(final_ocr_text) > 50000:
                        final_ocr_text = final_ocr_text[:50000] + "\n\n[... text trunchiat pentru a evita timeout-uri ...]"
                    text = final_ocr_text
                    print(f"✅ Text extras cu OCR din {len(images_to_process)} pagini (din {total_pages} totale): {len(text)} caractere")
                else:
                    return JSONResponse(
                        status_code=400,
                        content={"error": "Nu s-a putut extrage text din PDF nici cu OCR. PDF-ul poate fi protejat sau de calitate prea slabă."}
                    )
            except Exception as ocr_error:
                import traceback
                error_details = traceback.format_exc()
                error_str = str(ocr_error).lower()
                print(f"❌ Eroare la OCR pentru PDF: {error_details}")
                
                # Verifică dacă eroarea este legată de poppler
                if "poppler" in error_str or "pdftoppm" in error_str or "pdfinfo" in error_str:
                    return JSONResponse(
                        status_code=500,
                        content={
                            "error": "Poppler nu este instalat sau nu este în PATH. Pentru a procesa PDF-uri scanate, instalează poppler:\n"
                            "Windows: https://github.com/oschwartz10612/poppler-windows/releases\n"
                            "Linux: sudo apt-get install poppler-utils\n"
                            "macOS: brew install poppler"
                        }
                    )
                
                return JSONResponse(
                    status_code=500,
                    content={"error": f"Eroare la procesarea PDF cu OCR: {str(ocr_error)}"}
                )
        
        if not text.strip():
            return JSONResponse(
                status_code=400,
                content={"error": "Nu s-a putut extrage text din PDF. PDF-ul poate fi protejat sau de calitate prea slabă."}
            )
        
        # Limitează textul final la 50000 caractere pentru a evita timeout-uri
        final_text = text.strip()
        is_truncated = False
        if len(final_text) > 50000:
            final_text = final_text[:50000] + "\n\n[... text trunchiat pentru a evita timeout-uri ...]"
            is_truncated = True
        
        return JSONResponse(content={
            "text": final_text,
            "pages": len(pdf_reader.pages),
            "filename": pdf.filename,
            "method": "ocr" if not extracted_pages else "direct",
            "truncated": is_truncated,
            "original_length": len(text.strip())
        })
        
    except PyPDF2.errors.PdfReadError as e:
        return JSONResponse(
            status_code=400,
            content={"error": f"PDF corupt sau invalid: {str(e)}"}
        )
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"❌ Eroare detaliată la procesarea PDF: {error_details}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la procesarea PDF: {str(e)}. Verifică consola serverului pentru detalii."}
        )

@router.post("/extract-image")
async def extract_image(
    image: UploadFile = File(...),
    correct_text: bool = Form(False),
    expected_fields: Optional[str] = Form(None)
):
    """
    Extrage textul dintr-o imagine folosind OCR (PaddleOCR preferat, fallback la Tesseract).
    Opțional: corectează textul OCR și identifică datele lipsă.
    
    Args:
        image: Fișierul imagine
        correct_text: Dacă True, corectează automat textul OCR cu LLM
        expected_fields: Lista de câmpuri așteptate (JSON string sau comma-separated) pentru identificare date lipsă
    """
    print(f"📸 Primire cerere extragere text din imagine: {image.filename}, content_type: {image.content_type}")
    
    # Preferă PaddleOCR dacă este disponibil
    if PADDLEOCR_AVAILABLE_IMPORT:
        try:
            # Citește conținutul imaginii
            image_content = await image.read()
            if not image_content:
                return JSONResponse(
                    status_code=400,
                    content={"error": "Fișierul este gol sau nu a putut fi citit."}
                )
            
            print(f"📸 Procesare imagine cu PaddleOCR: {image.filename}, {len(image_content)} bytes")
            
            # Procesează cu PaddleOCR
            text, _ = process_image(image_content, lang='ro', return_boxes=False)
            
            if not text or not text.strip():
                print(f"⚠️ PaddleOCR nu a extras text, încerc cu Tesseract...")
                # Fallback la Tesseract
                if OCR_AVAILABLE:
                    # Continuă cu Tesseract mai jos
                    pass
                else:
                    return JSONResponse(
                        status_code=400,
                        content={"error": "Nu s-a putut extrage text din imagine. Imaginea poate să nu conțină text sau calitatea este prea slabă."}
                    )
            else:
                print(f"✅ Text extras cu PaddleOCR: {len(text)} caractere")
                extracted_text = text.strip()
                
                # Post-procesare: corectare text și identificare date lipsă
                response_data = {
                    "text": extracted_text,
                    "filename": image.filename,
                    "type": "image",
                    "method": "paddleocr"
                }
                
                if correct_text and correct_ocr_text:
                    print("🔧 Corectare text OCR cu LLM...")
                    try:
                        # Folosește modelul default (qwen2.5:7b) - va fi setat automat în postprocess.py
                        correction_result = correct_ocr_text(extracted_text, context=f"Document: {image.filename}")
                        response_data["corrected_text"] = correction_result.get("corrected_text", extracted_text)
                        response_data["corrections"] = correction_result.get("corrections", [])
                        response_data["correction_confidence"] = correction_result.get("confidence", 0.0)
                    except Exception as e:
                        print(f"⚠️ Eroare la corectarea textului OCR: {e}")
                        # Dacă corecția eșuează, folosește textul original
                        response_data["corrected_text"] = extracted_text
                        response_data["text"] = extracted_text  # Asigură-te că textul original este disponibil
                
                if expected_fields and identify_missing_fields:
                    import json
                    try:
                        if expected_fields.startswith('['):
                            fields_list = json.loads(expected_fields)
                        else:
                            fields_list = [f.strip() for f in expected_fields.split(',')]
                        
                        print(f"🔍 Identificare date lipsă pentru câmpuri: {fields_list}")
                        missing_result = identify_missing_fields(extracted_text, fields_list, context=f"Document: {image.filename}")
                        response_data["found_fields"] = missing_result.get("found_fields", [])
                        response_data["missing_fields"] = missing_result.get("missing_fields", [])
                        response_data["suggestions"] = missing_result.get("suggestions", "")
                    except Exception as e:
                        print(f"⚠️ Eroare la parsarea expected_fields sau identificarea datelor lipsă: {e}")
                
                return JSONResponse(content=response_data)
        except Exception as e:
            import traceback
            print(f"⚠️ Eroare la PaddleOCR: {traceback.format_exc()}")
            print(f"⚠️ Încerc cu Tesseract ca fallback...")
            # Fallback la Tesseract dacă PaddleOCR eșuează
            if not OCR_AVAILABLE:
                return JSONResponse(
                    status_code=500,
                    content={"error": f"PaddleOCR eșuat și Tesseract nu este disponibil. Eroare: {str(e)}"}
                )
            # Continuă cu Tesseract mai jos - resetează file pointer
            await image.seek(0)
    
    # Fallback la Tesseract
    if not OCR_AVAILABLE:
        print("❌ OCR nu este disponibil")
        return JSONResponse(
            status_code=500,
            content={"error": "OCR nu este disponibil. Rulează: pip install paddleocr opencv-python SAU pip install pytesseract pillow. Asigură-te că Tesseract OCR este instalat pe sistem."}
        )
    
    # Verifică tipul de fișier (verifică și extensia dacă content_type nu este setat)
    allowed_types = ["image/jpeg", "image/jpg", "image/png", "image/gif", "image/bmp", "image/webp", "image/x-png"]
    allowed_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
    
    # Normalizează content_type (unele browsere trimit "image/x-png" în loc de "image/png")
    normalized_content_type = image.content_type.lower() if image.content_type else ""
    if normalized_content_type == "image/x-png":
        normalized_content_type = "image/png"
    
    # Verifică content_type sau extensia fișierului
    is_valid_type = normalized_content_type in [t.lower() for t in allowed_types] if normalized_content_type else False
    is_valid_extension = any(image.filename.lower().endswith(ext) for ext in allowed_extensions) if image.filename else False
    
    print(f"🔍 Validare fișier: filename={image.filename}, content_type={image.content_type}, normalized={normalized_content_type}, is_valid_type={is_valid_type}, is_valid_extension={is_valid_extension}")
    
    # Dacă niciunul nu este valid, încercă să detecteze tipul din conținut
    if not is_valid_type and not is_valid_extension:
        print(f"⚠️ Tip de fișier nu este valid din header, verifică conținutul...")
        # Nu returnăm eroare imediat, vom încerca să deschidem imaginea și dacă reușește, continuăm
        # (PIL poate deschide imagini chiar dacă content_type nu este setat corect)
    
    try:
        # Citește conținutul imaginii
        try:
            image_content = await image.read()
            print(f"📦 Conținut citit: {len(image_content)} bytes")
        except Exception as read_error:
            import traceback
            print(f"❌ Eroare la citirea fișierului: {traceback.format_exc()}")
            return JSONResponse(
                status_code=400,
                content={"error": f"Eroare la citirea fișierului: {str(read_error)}"}
            )
        
        if not image_content:
            print("❌ Fișierul este gol")
            return JSONResponse(
                status_code=400,
                content={"error": "Fișierul este gol sau nu a putut fi citit."}
            )
        
        # Deschide imaginea cu PIL
        try:
            img = Image.open(io.BytesIO(image_content))
            print(f"✅ Imagine deschisă: {img.format}, {img.size}, mode: {img.mode}")
            
            # Convertește la RGB dacă este necesar (pentru PNG cu transparență sau alte formate)
            if img.mode != 'RGB':
                print(f"🔄 Conversie din {img.mode} la RGB")
                try:
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'RGBA':
                        rgb_img.paste(img, mask=img.split()[3] if len(img.split()) > 3 else None)  # Folosește canalul alpha ca mask
                    elif img.mode == 'P':
                        # Paletă de culori - convertește direct
                        rgb_img = img.convert('RGB')
                    else:
                        rgb_img.paste(img)
                    img = rgb_img
                    print(f"✅ Imagine convertită la RGB: {img.size}, mode: {img.mode}")
                except Exception as convert_error:
                    print(f"⚠️ Eroare la conversie, încercă convert direct: {convert_error}")
                    try:
                        img = img.convert('RGB')
                        print(f"✅ Imagine convertită direct la RGB")
                    except Exception as e2:
                        print(f"❌ Eroare la conversie: {e2}")
                        return JSONResponse(
                            status_code=400,
                            content={"error": f"Nu s-a putut converti imaginea la format RGB: {str(e2)}"}
                        )
        except Exception as e:
            import traceback
            print(f"❌ Eroare la deschiderea imaginii: {traceback.format_exc()}")
            return JSONResponse(
                status_code=400,
                content={"error": f"Nu s-a putut deschide imaginea. Verifică că fișierul este o imagine validă. Eroare: {str(e)}"}
            )
        
        # Extrage textul folosind OCR
        text = None
        error_msg = None
        
        # Verifică dacă Tesseract este disponibil încă o dată (în caz că OCR_AVAILABLE era True dar Tesseract nu funcționează)
        try:
            pytesseract.get_tesseract_version()
        except Exception as tess_check_error:
            error_msg = str(tess_check_error)
            if "tesseract" in error_msg.lower() or "not found" in error_msg.lower() or "no such file" in error_msg.lower():
                return JSONResponse(
                    status_code=500,
                    content={"error": f"Tesseract OCR nu este instalat sau nu este în PATH. Eroare: {error_msg}. Instalează Tesseract OCR de la: https://github.com/UB-Mannheim/tesseract/wiki"}
                )
        
        # Încearcă cu diferite configurații de limbi
        lang_configs = ['ron+eng', 'eng', 'ron', None]  # None = default
        
        print(f"🔤 Încearcă extragere text cu OCR...")
        text = None
        error_msg = None
        
        for lang_config in lang_configs:
            try:
                print(f"  📝 Încearcă cu limba: {lang_config or 'default'}")
                
                # Folosește timeout de 60 secunde pentru OCR
                try:
                    if lang_config:
                        text = pytesseract.image_to_string(img, lang=lang_config, timeout=60)
                    else:
                        text = pytesseract.image_to_string(img, timeout=60)
                except Exception as ocr_ex:
                    # Verifică dacă este timeout
                    error_str = str(ocr_ex).lower()
                    if "timeout" in error_str or "timed out" in error_str:
                        print(f"  ⏱️ Timeout la extragere OCR cu limba {lang_config or 'default'}")
                        error_msg = "Timeout la extragerea textului cu OCR (procesarea a durat prea mult)"
                        continue
                    else:
                        raise ocr_ex
                
                print(f"  ✅ Text extras: {len(text)} caractere (primul fragment: {text[:100] if text else 'N/A'})")
                
                if text and text.strip():
                    print(f"✅ Text extras cu succes cu limba: {lang_config or 'default'}")
                    break  # Dacă am obținut text, ieșim din loop
                else:
                    print(f"  ⚠️ Nu s-a extras text cu limba: {lang_config or 'default'}")
            except Exception as e:
                error_msg = str(e)
                print(f"  ❌ Eroare cu limba {lang_config or 'default'}: {error_msg}")
                # Dacă e eroare de Tesseract, oprește imediat
                if "tesseract" in error_msg.lower() or "not found" in error_msg.lower() or "no such file" in error_msg.lower():
                    print(f"❌ Tesseract nu este disponibil")
                    return JSONResponse(
                        status_code=500,
                        content={"error": f"Tesseract OCR nu este instalat sau nu este în PATH. Eroare: {error_msg}. Instalează Tesseract OCR de la: https://github.com/UB-Mannheim/tesseract/wiki"}
                    )
                # Continuă cu următoarea configurație pentru alte erori
                continue
        
        # Dacă nu am reușit să extragem text, verifică eroarea
        if not text or not text.strip():
            if error_msg:
                if "tesseract" in error_msg.lower() or "not found" in error_msg.lower() or "no such file" in error_msg.lower():
                    return JSONResponse(
                        status_code=500,
                        content={"error": f"Tesseract OCR nu este instalat sau nu este în PATH. Eroare: {error_msg}. Vezi INSTALARE_OCR.md pentru instrucțiuni."}
                    )
                else:
                    return JSONResponse(
                        status_code=500,
                        content={"error": f"Eroare la extragerea textului cu OCR: {error_msg}"}
                    )
            else:
                # Nu am eroare, dar nici text - probabil imaginea nu conține text
                return JSONResponse(
                    status_code=400,
                    content={"error": "Nu s-a putut extrage text din imagine. Imaginea poate să nu conțină text sau calitatea este prea slabă. Încearcă cu o imagine de calitate mai bună."}
                )
        
        extracted_text = text.strip()
        
        # Post-procesare: corectare text și identificare date lipsă
        response_data = {
            "text": extracted_text,
            "filename": image.filename,
            "type": "image",
            "method": "tesseract"
        }
        
        if correct_text and correct_ocr_text:
            print("🔧 Corectare text OCR cu LLM...")
            try:
                correction_result = correct_ocr_text(extracted_text, context=f"Document: {image.filename}")
                response_data["corrected_text"] = correction_result.get("corrected_text", extracted_text)
                response_data["corrections"] = correction_result.get("corrections", [])
                response_data["correction_confidence"] = correction_result.get("confidence", 0.0)
            except Exception as e:
                print(f"⚠️ Eroare la corectarea textului OCR: {e}")
                response_data["corrected_text"] = extracted_text
        
        if expected_fields and identify_missing_fields:
            import json
            try:
                if expected_fields.startswith('['):
                    fields_list = json.loads(expected_fields)
                else:
                    fields_list = [f.strip() for f in expected_fields.split(',')]
                
                print(f"🔍 Identificare date lipsă pentru câmpuri: {fields_list}")
                missing_result = identify_missing_fields(extracted_text, fields_list, context=f"Document: {image.filename}")
                response_data["found_fields"] = missing_result.get("found_fields", [])
                response_data["missing_fields"] = missing_result.get("missing_fields", [])
                response_data["suggestions"] = missing_result.get("suggestions", "")
            except Exception as e:
                print(f"⚠️ Eroare la parsarea expected_fields sau identificarea datelor lipsă: {e}")
        
        return JSONResponse(content=response_data)
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"❌ Eroare detaliată la procesarea imaginii: {error_details}")
        # Nu lăsăm serverul să se oprească - returnăm întotdeauna un răspuns
        try:
            return JSONResponse(
                status_code=500,
                content={"error": f"Eroare la procesarea imaginii: {str(e)}. Verifică consola serverului pentru detalii."}
            )
        except Exception as response_error:
            # Dacă chiar și returnarea răspunsului eșuează, loghează și re-raise
            print(f"❌❌ EROARE CRITICĂ: Nu s-a putut returna răspuns: {response_error}")
            print(f"❌❌ Eroare originală: {error_details}")
            raise

@router.post("/extract-docx")
async def extract_docx(
    docx: UploadFile = File(...)
):
    """
    Extrage textul dintr-un fișier DOCX.
    
    Args:
        docx: Fișierul DOCX
    """
    if not DOCX_AVAILABLE:
        return JSONResponse(
            status_code=500,
            content={"error": "python-docx nu este instalat. Rulează: pip install python-docx"}
        )
    
    # Verifică tipul de fișier
    allowed_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"  # Pentru compatibilitate cu .doc (vechi)
    ]
    allowed_extensions = ['.docx', '.doc']
    
    is_valid_type = docx.content_type in allowed_types if docx.content_type else False
    is_valid_extension = any(docx.filename.lower().endswith(ext) for ext in allowed_extensions) if docx.filename else False
    
    if not is_valid_type and not is_valid_extension:
        return JSONResponse(
            status_code=400,
            content={"error": "Fișierul trebuie să fie DOCX"}
        )
    
    try:
        # Citește conținutul fișierului
        docx_content = await docx.read()
        
        if not docx_content:
            return JSONResponse(
                status_code=400,
                content={"error": "Fișierul DOCX este gol sau nu a putut fi citit."}
            )
        
        # Deschide documentul DOCX
        doc = Document(io.BytesIO(docx_content))
        
        # Extrage textul din toate paragrafele
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extrage textul din tabele (dacă există)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            return JSONResponse(
                status_code=400,
                content={"error": "Nu s-a putut extrage text din DOCX. Documentul poate fi gol."}
            )
        
        # Combină toate părțile textului
        text = "\n".join(text_parts)
        
        # Limitează textul la 50000 caractere pentru a evita timeout-uri
        final_text = text.strip()
        is_truncated = False
        if len(final_text) > 50000:
            final_text = final_text[:50000] + "\n\n[... text trunchiat pentru a evita timeout-uri ...]"
            is_truncated = True
        
        return JSONResponse(content={
            "text": final_text,
            "filename": docx.filename,
            "type": "docx",
            "method": "python-docx",
            "truncated": is_truncated,
            "original_length": len(text.strip())
        })
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"❌ Eroare detaliată la procesarea DOCX: {error_details}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la procesarea DOCX: {str(e)}. Verifică consola serverului pentru detalii."}
        )

