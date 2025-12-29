from fastapi import APIRouter, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse, Response
from fastapi import Form, UploadFile, File
from typing import Optional
import os
import uuid
from models.schemas import ChatRequest
from database import create_client_chat, get_client_chat
from rag_manager import get_tenant_rag_store
from core.cache import get_cached_config
from core.conversation import get_tenant_id_from_chat_id
from core.config import PDF_AVAILABLE, OCR_AVAILABLE
from routers.chat import stream_response
from fastapi.responses import StreamingResponse
import PyPDF2

router = APIRouter(tags=["static"])

@router.get("/")
async def serve_index():
    return FileResponse("public/index.html")

@router.get("/scripts/{filename}")
async def serve_script(filename: str):
    file_path = os.path.join("public", "scripts", filename)
    if not os.path.exists(file_path):
        print(f"⚠️ Fișier script lipsă: {file_path}")
        return JSONResponse(status_code=404, content={"error": "File not found"})
    return FileResponse(file_path)

@router.get("/style/{filename}")
async def serve_style(filename: str):
    file_path = os.path.join("public", "style", filename)
    if not os.path.exists(file_path):
        print(f"⚠️ Fișier CSS lipsă: {file_path}")
        return JSONResponse(status_code=404, content={"error": "File not found"})
    return FileResponse(file_path)

@router.get("/favicon.ico")
async def favicon():
    # Return empty response pentru favicon (browserul va încerca să-l încarce)
    return Response(status_code=204)

@router.get("/evenimente", response_class=HTMLResponse)
async def serve_events():
    return FileResponse("public/formular_evenimente.html")

@router.get("/primarie", response_class=HTMLResponse)
async def serve_primarie():
    return FileResponse("public/site_primarie.html")

@router.get("/rezervari", response_class=HTMLResponse)
async def serve_reservations():
    return FileResponse("public/formular_rezervari.html")

@router.get("/login")
async def serve_login():
    """
    Redirecționează către pagina de login Next.js
    Dacă Next.js rulează pe același server, folosește redirect.
    Dacă rulează separat, ajustează URL-ul corespunzător.
    """
    # Construiește URL-ul pentru redirect din variabila de mediu
    nextjs_base_url = os.getenv('NEXTJS_URL', 'http://localhost:3000')
    nextjs_url = f"{nextjs_base_url}/login"
    return RedirectResponse(url=nextjs_url, status_code=307)

@router.get("/register")
async def serve_register():
    """
    Redirecționează către pagina de înregistrare Next.js
    """
    nextjs_base_url = os.getenv('NEXTJS_URL', 'http://localhost:3000')
    nextjs_url = f"{nextjs_base_url}/register"
    return RedirectResponse(url=nextjs_url, status_code=307)

@router.get("/builder", response_class=HTMLResponse)
async def builder_page():
    return FileResponse("public/builder.html")

@router.post("/builder/create")
async def create_chat(
    request: Request,
    name: str = Form(...),
    model: str = Form(...),
    prompt: str = Form(...),
    chat_title: Optional[str] = Form(None),
    chat_subtitle: Optional[str] = Form(None),
    chat_color: Optional[str] = Form("#3b82f6"),
    rag_files: Optional[list[UploadFile]] = File(None)
):
    # Generează ID unic
    chat_id = name.lower().replace(" ", "-") + "-" + str(uuid.uuid4())[:8]
    
    # Procesează fișierele RAG
    rag_content = []
    if rag_files:
        os.makedirs("rag", exist_ok=True)
        os.makedirs(f"rag/{chat_id}", exist_ok=True)
        for file in rag_files:
            if file.filename:
                # Salvează fișierul
                file_path = f"rag/{chat_id}/{file.filename}"
                with open(file_path, "wb") as f:
                    content = await file.read()
                    f.write(content)
                
                # Extrage text din fișier
                text_content = ""
                if file.filename.endswith('.pdf') and PDF_AVAILABLE:
                    try:
                        with open(file_path, "rb") as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            for page_num, page in enumerate(pdf_reader.pages):
                                try:
                                    page_text = page.extract_text()
                                    if page_text.strip():
                                        text_content += f"\n--- Pagina {page_num + 1} ---\n{page_text}\n"
                                except Exception as e:
                                    print(f"Eroare la extragerea paginii {page_num + 1} din {file.filename}: {e}")
                                    continue
                        
                        # Dacă nu s-a extras text (PDF scanat), încearcă OCR
                        if not text_content.strip() and OCR_AVAILABLE:
                            try:
                                print(f"PDF {file.filename} pare scanat, încerc OCR...")
                                # Pentru PDF-uri scanate, ar trebui să convertim paginile în imagini
                                # Pentru moment, doar logăm
                                print(f"⚠️ PDF {file.filename} nu conține text extractibil. Poate fi scanat.")
                            except Exception as e:
                                print(f"Eroare OCR pentru {file.filename}: {e}")
                    except Exception as e:
                        print(f"Eroare la extragerea textului din PDF {file.filename}: {e}")
                elif file.filename.endswith(('.txt', '.md')):
                    try:
                        with open(file_path, "r", encoding="utf-8") as f:
                            text_content = f.read()
                    except UnicodeDecodeError:
                        # Încearcă cu encoding diferit
                        try:
                            with open(file_path, "r", encoding="latin-1") as f:
                                text_content = f.read()
                        except Exception as e:
                            print(f"Eroare la citirea fișierului {file.filename}: {e}")
                    except Exception as e:
                        print(f"Eroare la citirea fișierului {file.filename}: {e}")
                elif file.filename.endswith(('.doc', '.docx')):
                    # Pentru DOC/DOCX, ar trebui python-docx
                    try:
                        try:
                            from docx import Document
                            doc = Document(file_path)
                            for para in doc.paragraphs:
                                if para.text.strip():
                                    text_content += para.text + "\n"
                            # Extrage și din tabele
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                    if row_text.strip():
                                        text_content += row_text + "\n"
                        except ImportError:
                            print(f"⚠️ python-docx nu este instalat. Pentru DOC/DOCX, rulează: pip install python-docx")
                            text_content = f"[Fișier DOC/DOCX - instalează python-docx pentru extragere: pip install python-docx]"
                    except Exception as e:
                        print(f"Eroare la extragerea textului din DOC/DOCX {file.filename}: {e}")
                
                if text_content and text_content.strip():
                    # Nu limităm aici - limităm doar în prompt pentru a păstra tot conținutul în config
                    rag_content.append({
                        "filename": file.filename,
                        "content": text_content.strip()  # Păstrăm tot conținutul în config
                    })
                    print(f"✅ Text extras din {file.filename}: {len(text_content)} caractere")
                else:
                    print(f"⚠️ Nu s-a putut extrage text din {file.filename} (poate fi gol, scanat sau protejat)")
    
    # IMPORTANT: Nu modificăm prompt-ul salvat în config
    # Îmbunătățirea se face dinamic la runtime în funcție de context
    
    # Creează chatbot-ul în baza de date
    client_chat_id = create_client_chat(
        name=name,
        model=model,
        prompt=prompt,
        chat_title=chat_title or name,
        chat_subtitle=chat_subtitle or "Asistentul tău inteligent pentru găsirea informațiilor",
        chat_color=chat_color or "#3b82f6"
    )
    
    if not client_chat_id:
        return JSONResponse(
            status_code=500,
            content={"error": "Eroare la crearea chatbot-ului în baza de date"}
        )
    
    # Adaugă fișierele RAG în baza de date și vector store
    tenant_id = str(client_chat_id)
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        # Șterge toate documentele vechi
        rag_store.clear()
        
        # Adaugă documentele noi
        for item in rag_content:
            # Salvează în DB cu conținutul
            from database import add_rag_file
            add_rag_file(client_chat_id, item["filename"], item["content"])
            # Adaugă în vector store
            rag_store.add_document(item["filename"], item["content"])
        
        print(f"✅ Vector store creat pentru tenant {tenant_id}")
    except Exception as e:
        print(f"⚠️ Eroare la crearea vector store pentru tenant {tenant_id}: {e}")
    
    # Reîncarcă config-ul din DB
    config = get_cached_config(str(client_chat_id))

    # Returnează link direct pentru chat full-page
    chat_url = f"/chat/{client_chat_id}"
    
    return JSONResponse(content={
        "chat_id": str(client_chat_id),
        "chat_url": chat_url,
        "config": config
    })

@router.post("/ask")
async def ask_default(request: ChatRequest):
    # Obține config-ul din baza de date dacă există chat_id
    config = None
    model = "qwen2.5:7b"  # Default fallback
    prompt = "Ești asistentul Integra AI. Răspunde clar și politicos."
    
    if request.chat_id:
        # Obține config-ul din cache/baza de date
        config = get_cached_config(request.chat_id)
        if config:
            model = config.get("model", model)
            prompt = config.get("prompt", prompt)
    else:
        # Dacă nu există chat_id, folosește config default
        # Poți adăuga logică pentru a obține primul chat disponibil sau un chat default
        print("⚠️ Nu s-a furnizat chat_id în request, folosesc config default")

    messages = [
        {"role": "system", "content": prompt},
        {"role": "user", "content": request.message}
    ]

    # Extrage conținutul RAG din config dacă există
    rag_content = config.get("rag_content", []) if config else []
    institution_data = config.get("institution") if config else None
    tenant_id = config.get("tenant_id") if config else None

    return StreamingResponse(
        stream_response(messages, model, request.page_context, request.pdf_text, rag_content, institution_data, None, tenant_id),
        media_type="text/plain; charset=utf-8"
    )

