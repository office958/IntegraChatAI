from fastapi import APIRouter, HTTPException, status, UploadFile, File
from fastapi.responses import JSONResponse
import os
from urllib.parse import unquote
from database import (
    get_client_chat, create_client_chat, update_client_chat, list_all_client_chats,
    create_or_update_client_type,
    add_rag_file, delete_rag_file
)
from rag_manager import get_tenant_rag_store
from core.cache import get_cached_config, invalidate_config_cache
from core.conversation import get_tenant_id_from_chat_id
from core.config import PDF_AVAILABLE, OCR_AVAILABLE
import PyPDF2

router = APIRouter(prefix="/admin", tags=["admin"])

@router.put("/tenant/{chat_id}/institution")
async def update_institution_data(chat_id: str, institution_data: dict):
    """Actualizează datele instituției pentru un tenant"""
    # Convertește chat_id la int
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        # Caută după name
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Salvează în baza de date
    success = create_or_update_client_type(
        client_chat_id=client_chat_id,
        name=institution_data.get("name", ""),
        type=institution_data.get("type", "alta"),
        address=institution_data.get("address"),
        phone=institution_data.get("phone"),
        email=institution_data.get("email"),
        website=institution_data.get("website")
    )
    
    if not success:
        return JSONResponse(
            status_code=500,
            content={"error": "Eroare la actualizarea datelor instituției"}
        )
    
    # Invalidează cache-ul
    invalidate_config_cache(chat_id)
    
    # Reîncarcă config-ul
    config = get_cached_config(chat_id)
    
    return JSONResponse(content={
        "success": True,
        "message": "Datele instituției au fost actualizate",
        "config": config
    })

@router.put("/tenant/{chat_id}/config")
async def update_tenant_config(chat_id: str, config_updates: dict):
    """Actualizează configurația unui tenant (prompt, model, setări UI, etc.)"""
    # Convertește chat_id la int
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        # Caută după name
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Actualizează în baza de date
    success = update_client_chat(
        chat_id=client_chat_id,
        name=config_updates.get("name"),
        model=config_updates.get("model"),
        prompt=config_updates.get("prompt"),
        chat_title=config_updates.get("chat_title"),
        chat_subtitle=config_updates.get("chat_subtitle"),
        chat_color=config_updates.get("chat_color"),
        is_active=config_updates.get("is_active")
    )
    
    if not success:
        return JSONResponse(
            status_code=500,
            content={"error": "Eroare la actualizarea configurației"}
        )
    
    # Invalidează cache-ul
    invalidate_config_cache(chat_id)
    
    # Reîncarcă config-ul
    config = get_cached_config(chat_id)
    
    return JSONResponse(content={
        "success": True,
        "message": "Configurația a fost actualizată",
        "config": config
    })

@router.post("/tenant/{chat_id}/rag/upload")
async def upload_rag_file(chat_id: str, file: UploadFile = File(...)):
    """Încarcă un fișier RAG pentru un tenant"""
    print(f"📤 Upload RAG pentru tenant {chat_id}, fișier: {file.filename if file.filename else 'N/A'}")
    
    config = get_cached_config(chat_id)
    if not config:
        print(f"❌ Config nu există pentru {chat_id}")
        return JSONResponse(
            status_code=404,
            content={"error": f"Chat configuration not found: {chat_id}"}
        )
    
    if not file.filename:
        print(f"❌ Fișier fără nume pentru {chat_id}")
        return JSONResponse(
            status_code=400,
            content={"error": "Fișierul trebuie să aibă un nume"}
        )
    
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    
    # Citește fișierul în memorie (nu pe disk)
    try:
        file_data = await file.read()
        print(f"✅ Fișier citit: {len(file_data)} bytes")
    except Exception as e:
        print(f"❌ Eroare la citirea fișierului: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la citirea fișierului: {str(e)}"}
        )
    
    # Extrage text din fișierul din memorie
    text_content = ""
    try:
        print(f"📄 Încep extragerea textului din {file.filename}...")
        if file.filename.endswith('.pdf') and PDF_AVAILABLE:
            print(f"📄 Procesare PDF: {file.filename}")
            import io
            pdf_file = io.BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            print(f"📄 PDF are {len(pdf_reader.pages)} pagini")
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += f"\n--- Pagina {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    print(f"⚠️ Eroare la extragerea paginii {page_num + 1} din {file.filename}: {e}")
            if not text_content.strip():
                print(f"⚠️ PDF {file.filename} nu conține text extractibil (poate fi scanat)")
        elif file.filename.endswith(('.txt', '.md')):
            print(f"📄 Procesare text: {file.filename}")
            try:
                text_content = file_data.decode('utf-8')
            except UnicodeDecodeError:
                print(f"⚠️ Eroare encoding UTF-8, încerc latin-1...")
                text_content = file_data.decode('latin-1')
        elif file.filename.endswith(('.doc', '.docx')):
            print(f"📄 Procesare DOC/DOCX: {file.filename}")
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(file_data))
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content += para.text + "\n"
                # Extrage și din tabele
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_content += row_text + "\n"
            except ImportError:
                print(f"⚠️ python-docx nu este instalat. Pentru DOC/DOCX, rulează: pip install python-docx")
                text_content = f"[Fișier DOC/DOCX - instalează python-docx pentru extragere]"
        else:
            print(f"⚠️ Tip de fișier necunoscut: {file.filename}")
        
        print(f"✅ Text extras: {len(text_content)} caractere")
    except Exception as e:
        print(f"❌ Eroare la extragerea textului din {file.filename}: {e}")
        import traceback
        traceback.print_exc()
    
    # Convertește chat_id la int pentru DB
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Adaugă fișierul în baza de date cu conținutul text și fișierul binar
    if text_content and text_content.strip():
        # Salvează sau actualizează fișierul în DB cu conținutul text și fișierul binar
        add_rag_file(client_chat_id, file.filename, text_content.strip(), file_data)
        print(f"✅ Fișier RAG salvat în DB cu conținut și date: {file.filename} ({len(text_content)} caractere, {len(file_data)} bytes)")
    else:
        # Dacă nu s-a putut extrage text, salvează doar fișierul binar
        add_rag_file(client_chat_id, file.filename, None, file_data)
        print(f"⚠️ Nu s-a putut extrage text din {file.filename} (poate fi gol, scanat sau protejat) - salvat doar fișierul binar ({len(file_data)} bytes)")
    
    # Actualizează vector store
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        if text_content and text_content.strip():
            rag_store.add_document(file.filename, text_content.strip())
            print(f"✅ Fișier RAG adăugat în vector store pentru tenant {tenant_id}")
        else:
            print(f"⚠️ Nu s-a adăugat în vector store (fără conținut text)")
    except Exception as e:
        print(f"⚠️ Eroare la actualizarea vector store pentru tenant {tenant_id}: {e}")
        import traceback
        traceback.print_exc()
    
    # Invalidează cache-ul
    invalidate_config_cache(chat_id)
    
    return JSONResponse(content={
        "success": True,
        "message": f"Fișier {file.filename} încărcat cu succes" + ("" if text_content and text_content.strip() else " (fără conținut text extractibil)"),
        "filename": file.filename,
        "has_content": bool(text_content and text_content.strip()),
        "content_length": len(text_content) if text_content else 0
    })

@router.delete("/tenant/{chat_id}/rag/{filename}")
async def delete_rag_file_endpoint(chat_id: str, filename: str):
    """Șterge un fișier RAG pentru un tenant"""
    # Decode filename dacă este URL-encoded
    filename = unquote(filename)
    
    print(f"🗑️ Ștergere RAG pentru tenant {chat_id}, fișier: {filename}")
    
    config = get_cached_config(chat_id)
    if not config:
        print(f"❌ Config nu există pentru {chat_id}")
        return JSONResponse(
            status_code=404,
            content={"error": f"Chat configuration not found: {chat_id}"}
        )
    
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    
    # Nu mai ștergem fișierul de pe disk - doar din baza de date
    # (Fișierele sunt acum stocate în baza de date)
    
    # Convertește chat_id la int pentru DB
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Șterge din baza de date
    deleted = delete_rag_file(client_chat_id, filename)
    if deleted:
        print(f"✅ Fișier șters din DB: {filename}")
    else:
        print(f"⚠️ Fișier nu era în DB: {filename}")
    
    # Actualizează vector store
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        rag_store.remove_document(filename)
        print(f"✅ Fișier RAG șters din vector store pentru tenant {tenant_id}")
    except Exception as e:
        print(f"⚠️ Eroare la actualizarea vector store pentru tenant {tenant_id}: {e}")
        import traceback
        traceback.print_exc()
    
    # Invalidează cache-ul
    invalidate_config_cache(chat_id)
    
    return JSONResponse(content={
        "success": True,
        "message": f"Fișier {filename} șters cu succes"
    })

@router.get("/tenants")
async def list_all_tenants():
    """Listează toți tenant-ii (pentru panoul de administrare)"""
    # Încarcă din baza de date
    db_tenants = list_all_client_chats()
    
    tenants = []
    for db_tenant in db_tenants:
        tenants.append({
            "id": str(db_tenant.get("id", "")),
            "tenant_id": str(db_tenant.get("id", "")),
            "name": db_tenant.get("name", "Unknown"),
            "model": db_tenant.get("model", "unknown"),
            "is_active": bool(db_tenant.get("is_active", True)),
            "created_at": db_tenant.get("created_at"),
            "updated_at": db_tenant.get("updated_at"),
            "institution": db_tenant.get("institution"),
            "rag_files_count": db_tenant.get("rag_files_count", 0),
            "chat_title": db_tenant.get("chat_title"),
            "chat_color": db_tenant.get("chat_color")
        })
    
    return JSONResponse(content={"tenants": tenants})

@router.post("/tenant/create")
async def create_tenant(request: dict):
    """Creează un nou tenant/client chatbot"""
    try:
        name = request.get("name", "Chat nou")
        model = request.get("model", "qwen2.5:7b")
        prompt = request.get("prompt", "Ești asistentul Integra AI. Răspunde clar și politicos la întrebările utilizatorilor.")
        chat_title = request.get("chat_title", name)
        chat_subtitle = request.get("chat_subtitle", "Asistentul tău inteligent pentru găsirea informațiilor")
        chat_color = request.get("chat_color", "#3b82f6")
        
        # Creează chatbot-ul în baza de date
        client_chat_id = create_client_chat(
            name=name,
            model=model,
            prompt=prompt,
            chat_title=chat_title,
            chat_subtitle=chat_subtitle,
            chat_color=chat_color
        )
        
        if not client_chat_id:
            return JSONResponse(
                status_code=500,
                content={"error": "Eroare la crearea chatbot-ului în baza de date"}
            )
        
        # Reîncarcă config-ul din DB
        config = get_cached_config(str(client_chat_id))
        
        return JSONResponse(content={
            "success": True,
            "message": "Client creat cu succes",
            "tenant": {
                "id": str(client_chat_id),
                "tenant_id": str(client_chat_id),
                "name": name,
                "model": model,
                "is_active": True,
                "institution": None,
                "rag_files_count": 0,
                "chat_title": chat_title,
                "chat_color": chat_color
            }
        })
    except Exception as e:
        print(f"❌ Eroare la crearea tenant-ului: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la crearea clientului: {str(e)}"}
        )

@router.post("/tenant/{chat_id}/reprocess-rag")
async def reprocess_rag(chat_id: str):
    """Re-procesează fișierele RAG pentru un chat existent"""
    config = get_cached_config(chat_id)
    if not config:
        return JSONResponse(
            status_code=404,
            content={"error": f"Chat configuration not found: {chat_id}"}
        )
    
    rag_dir = f"rag/{chat_id}"
    
    if not os.path.exists(rag_dir):
        return JSONResponse(
            status_code=404,
            content={"error": f"Directorul RAG nu există: {rag_dir}"}
        )
    
    # Dacă rag_files este gol, detectează automat fișierele din director
    rag_files = config.get("rag_files", [])
    if not rag_files:
        # Detectează automat toate fișierele din director
        if os.path.exists(rag_dir):
            rag_files = [f for f in os.listdir(rag_dir) if os.path.isfile(os.path.join(rag_dir, f))]
            print(f"✅ Detectate automat {len(rag_files)} fișiere RAG în {rag_dir}")
    
    if not rag_files:
        return JSONResponse(
            status_code=400,
            content={"error": "Nu există fișiere RAG pentru acest chat"}
        )
    
    rag_content = []
    
    for filename in rag_files:
        file_path = os.path.join(rag_dir, filename)
        if not os.path.exists(file_path):
            print(f"⚠️ Fișier RAG nu există: {file_path}")
            continue
        
        text_content = ""
        try:
            if filename.endswith('.pdf') and PDF_AVAILABLE:
                with open(file_path, "rb") as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content += f"\n--- Pagina {page_num + 1} ---\n{page_text}\n"
                        except Exception as e:
                            print(f"Eroare la extragerea paginii {page_num + 1} din {filename}: {e}")
                            continue
                
                # Dacă nu s-a extras text (PDF scanat), logăm
                if not text_content.strip() and OCR_AVAILABLE:
                    print(f"⚠️ PDF {filename} pare scanat sau nu conține text extractibil.")
            elif filename.endswith(('.txt', '.md')):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text_content = f.read()
                except UnicodeDecodeError:
                    try:
                        with open(file_path, "r", encoding="latin-1") as f:
                            text_content = f.read()
                    except Exception as e:
                        print(f"Eroare la citirea fișierului {filename}: {e}")
                except Exception as e:
                    print(f"Eroare la citirea fișierului {filename}: {e}")
            elif filename.endswith(('.doc', '.docx')):
                try:
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text_content += para.text + "\n"
                        # Extrage și din tabele
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                if row_text.strip():
                                    text_content += row_text + "\n"
                    except ImportError:
                        print(f"⚠️ python-docx nu este instalat. Pentru DOC/DOCX, rulează: pip install python-docx")
                        text_content = f"[Fișier DOC/DOCX - instalează python-docx pentru extragere: pip install python-docx]"
                except Exception as e:
                    print(f"Eroare la extragerea textului din DOC/DOCX {filename}: {e}")
        except Exception as e:
            print(f"Eroare la procesarea {filename}: {e}")
        
        if text_content and text_content.strip():
            rag_content.append({
                "filename": filename,
                "content": text_content.strip()
            })
            print(f"✅ Text re-extras din {filename}: {len(text_content)} caractere")
    
    # Actualizează vector store-ul pentru tenant
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        # Șterge toate documentele vechi
        rag_store.clear()
        # Adaugă documentele noi
        for item in rag_content:
            rag_store.add_document(item["filename"], item["content"])
        print(f"✅ Vector store actualizat pentru tenant {tenant_id}")
    except Exception as e:
        print(f"⚠️ Eroare la actualizarea vector store pentru tenant {tenant_id}: {e}")
    
    # Invalidează cache-ul
    invalidate_config_cache(chat_id)
    
    return JSONResponse(content={
        "success": True,
        "processed_files": len(rag_content),
        "total_files": len(rag_files),
        "message": f"Re-procesat {len(rag_content)} din {len(rag_files)} fișiere și actualizat vector store"
    })

