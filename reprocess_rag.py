#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pentru re-procesarea fișierelor RAG pentru un chat existent
"""
import json
import os
import sys

# Setează encoding pentru stdout (Windows)
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# Adaugă directorul curent la path pentru a putea importa din main.py
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[WARNING] PyPDF2 nu este instalat. Ruleaza: pip install PyPDF2")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def reprocess_rag(chat_id: str):
    """Re-procesează fișierele RAG pentru un chat"""
    # Folosește path-uri absolute bazate pe directorul scriptului
    script_dir = os.path.dirname(os.path.abspath(__file__))
    configs_dir = os.path.join(script_dir, "configs")
    
    # Caută fișierul config (poate avea encoding diferit în nume)
    # Extrage hash-ul din chat_id (ultimele 8 caractere după ultimul -)
    chat_hash = chat_id.split('-')[-1] if '-' in chat_id else chat_id[-8:]
    config_path = None
    if os.path.exists(configs_dir):
        for filename in os.listdir(configs_dir):
            if filename.endswith('.json') and chat_hash in filename:
                config_path = os.path.join(configs_dir, filename)
                print(f"[INFO] Gasit config: {filename}")
                break
    
    if not config_path or not os.path.exists(config_path):
        print(f"[ERROR] Config nu exista pentru chat_id: {chat_id}")
        print(f"[INFO] Cauta in: {configs_dir}")
        return False
    
    # Încarcă config-ul
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    
    # Extrage chat_id real din numele fișierului config
    actual_chat_id = os.path.splitext(os.path.basename(config_path))[0]
    rag_dir = os.path.join(script_dir, "rag", actual_chat_id)
    
    if not os.path.exists(rag_dir):
        print(f"[ERROR] Directorul RAG nu exista: {rag_dir}")
        # Încearcă să găsească directorul RAG după hash
        rag_base = os.path.join(script_dir, "rag")
        if os.path.exists(rag_base):
            for dirname in os.listdir(rag_base):
                if chat_hash in dirname:
                    rag_dir = os.path.join(rag_base, dirname)
                    print(f"[INFO] Gasit director RAG: {dirname}")
                    break
        if not os.path.exists(rag_dir):
            print(f"[ERROR] Nu s-a gasit directorul RAG pentru chat_id: {chat_id}")
            return False
    
    # Dacă rag_files este gol, detectează automat fișierele din director
    rag_files = config.get("rag_files", [])
    if not rag_files:
        rag_files = [f for f in os.listdir(rag_dir) if os.path.isfile(os.path.join(rag_dir, f))]
        config["rag_files"] = rag_files
        print(f"[OK] Detectate automat {len(rag_files)} fisiere RAG")
    
    if not rag_files:
        print("[ERROR] Nu exista fisiere RAG")
        return False
    
    rag_content = []
    
    for filename in rag_files:
        file_path = os.path.join(rag_dir, filename)
        if not os.path.exists(file_path):
            print(f"[WARNING] Fisier nu exista: {file_path}")
            continue
        
        text_content = ""
        try:
            if filename.endswith('.pdf') and PDF_AVAILABLE:
                with open(file_path, "rb") as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content += f"\n--- Pagina {page_num + 1} ---\n{page_text}\n"
                        except Exception as e:
                            print(f"[WARNING] Eroare la pagina {page_num + 1} din {filename}: {e}")
                            continue
                
                if not text_content.strip():
                    print(f"[WARNING] PDF {filename} nu contine text extractibil (poate fi scanat)")
            elif filename.endswith(('.txt', '.md')):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text_content = f.read()
                except UnicodeDecodeError:
                    try:
                        with open(file_path, "r", encoding="latin-1") as f:
                            text_content = f.read()
                    except Exception as e:
                        print(f"[WARNING] Eroare la citirea {filename}: {e}")
                except Exception as e:
                    print(f"[WARNING] Eroare la citirea {filename}: {e}")
            elif filename.endswith(('.doc', '.docx')):
                if DOCX_AVAILABLE:
                    try:
                        doc = Document(file_path)
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text_content += para.text + "\n"
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                if row_text.strip():
                                    text_content += row_text + "\n"
                    except Exception as e:
                        print(f"[WARNING] Eroare la extragerea din DOC/DOCX {filename}: {e}")
                else:
                    print(f"[WARNING] python-docx nu este instalat pentru {filename}")
        except Exception as e:
            print(f"[WARNING] Eroare la procesarea {filename}: {e}")
        
        if text_content and text_content.strip():
            rag_content.append({
                "filename": filename,
                "content": text_content.strip()
            })
            print(f"[OK] Text extras din {filename}: {len(text_content)} caractere")
        else:
            print(f"[WARNING] Nu s-a putut extrage text din {filename}")
    
    # Actualizează config-ul
    config["rag_content"] = rag_content
    with open(config_path, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    
    # Invalidează cache-ul dacă există (pentru main.py)
    print(f"[INFO] Config actualizat: {config_path}")
    
    print(f"\n[OK] Re-procesare completa: {len(rag_content)} din {len(rag_files)} fisiere procesate cu succes")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Utilizare: python reprocess_rag.py <chat_id>")
        print("Exemplu: python reprocess_rag.py asistentul-virtual-al-primăriei-tășnad-dfe779d4")
        sys.exit(1)
    
    chat_id = sys.argv[1]
    success = reprocess_rag(chat_id)
    sys.exit(0 if success else 1)

