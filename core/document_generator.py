"""
Generator de documente DOCX și PDF din template-uri HTML cu variabile
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
from typing import Dict, List
from datetime import datetime
import html

def clean_html_text(html_content: str) -> str:
    """Elimină tag-urile HTML și păstrează doar textul"""
    # Elimină tag-urile HTML
    text = re.sub(r'<[^>]+>', '', html_content)
    # Decodifică entitățile HTML
    text = html.unescape(text)
    return text.strip()

def parse_html_to_paragraphs(html_content: str) -> List[Dict]:
    """Parsează HTML simplu și returnează o listă de paragrafe cu formatare"""
    paragraphs = []
    
    # Extrage paragrafele (separate de <p>, <br>, sau newlines)
    parts = re.split(r'<p[^>]*>|</p>|<br\s*/?>|\n\n+', html_content)
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
        
        # Elimină tag-urile HTML rămase
        text = clean_html_text(part)
        if text:
            # Detectează formatare simplă
            is_bold = '<strong>' in part or '<b>' in part
            is_italic = '<em>' in part or '<i>' in part
            is_underline = '<u>' in part
            
            # Detectează header-uri
            header_level = None
            if re.search(r'<h1[^>]*>', part):
                header_level = 1
            elif re.search(r'<h2[^>]*>', part):
                header_level = 2
            elif re.search(r'<h3[^>]*>', part):
                header_level = 3
            
            paragraphs.append({
                'text': text,
                'bold': is_bold,
                'italic': is_italic,
                'underline': is_underline,
                'header_level': header_level
            })
    
    return paragraphs

def generate_docx_from_template(
    html_content: str,
    output_filename: str,
    variables: Dict[str, str] = None
) -> str:
    """Generează un fișier DOCX din HTML template cu variabile înlocuite"""
    
    # Înlocuiește variabilele
    if variables:
        for key, value in variables.items():
            placeholder = f'{{{{ ${key} }}}}'
            html_content = html_content.replace(placeholder, str(value))
    
    # Înlocuiește data curentă dacă nu este setată
    if '{{ $today }}' in html_content:
        today = datetime.now().strftime('%d.%m.%Y')
        html_content = html_content.replace('{{ $today }}', today)
    
    # Parsează HTML în paragrafe
    paragraphs_data = parse_html_to_paragraphs(html_content)
    
    # Creează documentul DOCX
    doc = Document()
    
    # Setează fontul default
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for para_data in paragraphs_data:
        if para_data['header_level']:
            # Adaugă header
            heading_style = f'Heading {para_data["header_level"]}'
            para = doc.add_heading(para_data['text'], level=para_data['header_level'])
        else:
            # Adaugă paragraf normal
            para = doc.add_paragraph()
            run = para.add_run(para_data['text'])
            
            # Aplică formatare
            if para_data['bold']:
                run.bold = True
            if para_data['italic']:
                run.italic = True
            if para_data['underline']:
                run.underline = True
        
        # Setează spațierea
        para_format = para.paragraph_format
        para_format.space_after = Pt(6)
        para_format.space_before = Pt(0)
    
    # Salvează documentul
    output_dir = "pdf_generated"
    os.makedirs(output_dir, exist_ok=True)
    
    # Generează nume unic pentru fișier
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    base_name = os.path.splitext(output_filename)[0]
    output_path = os.path.join(output_dir, f"{base_name}_{timestamp}.docx")
    
    doc.save(output_path)
    print(f"✅ Document DOCX generat: {output_path}")
    
    return output_path

def convert_docx_to_pdf(docx_path: str) -> str:
    """Convertește DOCX în PDF"""
    try:
        # Încearcă cu docx2pdf (Windows)
        try:
            from docx2pdf import convert
            pdf_path = docx_path.replace('.docx', '.pdf')
            convert(docx_path, pdf_path)
            print(f"✅ Document PDF generat: {pdf_path}")
            return pdf_path
        except ImportError:
            pass
        
        # Fallback: folosește LibreOffice (dacă este instalat)
        import subprocess
        pdf_path = docx_path.replace('.docx', '.pdf')
        try:
            subprocess.run([
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path),
                docx_path
            ], check=True, capture_output=True)
            print(f"✅ Document PDF generat cu LibreOffice: {pdf_path}")
            return pdf_path
        except (subprocess.CalledProcessError, FileNotFoundError):
            pass
        
        # Dacă nu există niciun convertor, returnează None
        print("⚠️ Nu s-a putut converti DOCX în PDF. Instalează docx2pdf sau LibreOffice.")
        return None
        
    except Exception as e:
        print(f"❌ Eroare la conversia DOCX în PDF: {e}")
        return None

def generate_document_from_template(
    template_html: str,
    filename: str,
    variables: Dict[str, str] = None
) -> Dict[str, str]:
    """Generează atât DOCX cât și PDF din template"""
    
    # Generează DOCX
    docx_path = generate_docx_from_template(template_html, filename, variables)
    
    # Convertește în PDF
    pdf_path = convert_docx_to_pdf(docx_path)
    
    return {
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "docx_url": f"/pdf_generated/{os.path.basename(docx_path)}",
        "pdf_url": f"/pdf_generated/{os.path.basename(pdf_path)}" if pdf_path else None
    }



