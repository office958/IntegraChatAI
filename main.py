from fastapi import FastAPI, Form, Request, UploadFile, File, Depends, HTTPException, status
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse, JSONResponse, Response, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr
from ollama import Client
import json, os, asyncio, uuid, re, io
from functools import lru_cache
from typing import Optional
from datetime import datetime, timedelta
import bcrypt
import jwt

# Importă modulele noastre pentru RAG și prompt building
from rag_manager import get_tenant_rag_store
from prompt_builder import build_dynamic_system_prompt
from database import (
    get_client_chat, create_client_chat, update_client_chat, list_all_client_chats,
    get_client_type, create_or_update_client_type,
    get_rag_files, add_rag_file, delete_rag_file,
    create_chat_session, get_chat_session, list_user_chat_sessions, update_chat_session as db_update_chat_session, delete_chat_session as db_delete_chat_session,
    get_conversation_history as db_get_conversation_history,
    add_message_to_conversation as db_add_message_to_conversation,
    clear_conversation_history as db_clear_conversation_history,
    get_user, create_user
)

# Încarcă variabilele de mediu din .env
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    print("⚠️ python-dotenv nu este instalat. Pentru a folosi .env, rulează: pip install python-dotenv")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ PyPDF2 nu este instalat. Rulează: pip install PyPDF2")

try:
    import pytesseract
    from PIL import Image
    import platform
    
    # Verifică dacă Tesseract este disponibil
    try:
        pytesseract.get_tesseract_version()
        OCR_AVAILABLE = True
        print("✅ OCR disponibil - Tesseract funcționează")
    except Exception as tess_error:
        # Încearcă să configureze calea Tesseract (doar pe Windows, dacă nu e în PATH)
        if platform.system() == 'Windows':
            possible_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            ]
            tesseract_found = False
            for path in possible_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    try:
                        pytesseract.get_tesseract_version()
                        print(f"✅ Tesseract găsit și configurat la: {path}")
                        OCR_AVAILABLE = True
                        tesseract_found = True
                        break
                    except:
                        continue
            
            if not tesseract_found:
                OCR_AVAILABLE = False
                print(f"⚠️ Tesseract nu este disponibil. Eroare: {str(tess_error)}")
                print("💡 Instalează Tesseract OCR de la: https://github.com/UB-Mannheim/tesseract/wiki")
        else:
            OCR_AVAILABLE = False
            print(f"⚠️ Tesseract nu este disponibil. Eroare: {str(tess_error)}")
            print("💡 Instalează Tesseract OCR: sudo apt-get install tesseract-ocr (Linux) sau brew install tesseract (macOS)")
            
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR nu este disponibil. Rulează: pip install pytesseract pillow")
except Exception as e:
    OCR_AVAILABLE = False
    print(f"⚠️ OCR nu este disponibil. Eroare: {str(e)}")

app = FastAPI(title="Integra AI Builder")

# Conectare la Ollama - citeste IP-ul din variabilele de mediu
OLLAMA_HOST = os.getenv('OLLAMA_HOST', 'localhost:11434')
ollama = Client(host=OLLAMA_HOST)

# JWT Configuration
JWT_SECRET_KEY = os.getenv('JWT_SECRET_KEY', 'your-secret-key-change-in-production')
JWT_ALGORITHM = os.getenv('JWT_ALGORITHM', 'HS256')
JWT_EXPIRATION_HOURS = int(os.getenv('JWT_EXPIRATION_HOURS', str(24 * 7)))  # Default: 7 zile

# Security - folosit pentru endpoint-uri care necesită autentificare obligatorie
security = HTTPBearer()
# Security optional - pentru endpoint-uri care permit și request-uri fără autentificare
security_optional = HTTPBearer(auto_error=False)

# Cache pentru config-uri (se reîncarcă automat când se modifică)
_config_cache = {}
_config_cache_timestamps = {}

# Stocare istoric conversații (keyed by tenant_id/chat_id)
# Format: {tenant_id: {chat_id: [{"role": "user/assistant", "content": "..."}, ...]}}
# Asigură izolarea completă între tenant-i
_conversation_history = {}

# Funcție helper pentru a obține tenant_id din chat_id
def get_tenant_id_from_chat_id(chat_id: str) -> str:
    """
    Extrage tenant_id din chat_id.
    În structura actuală, chat_id este deja identificatorul tenant-ului.
    Pentru compatibilitate, folosim chat_id ca tenant_id.
    """
    return chat_id

# Dimensiune maximă context window (în tokens aproximativi, folosim caractere ca proxy)
# Pentru majoritatea modelelor, ~4 caractere = 1 token
MAX_CONTEXT_CHARS = int(os.getenv('MAX_CONTEXT_CHARS', '32000'))  # ~8000 tokens (ajustabil în funcție de model)
CONTEXT_RESERVE = int(os.getenv('CONTEXT_RESERVE', '2000'))  # Rezervă pentru system prompt și mesajul curent

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# === Modele pentru autentificare ===
class LoginRequest(BaseModel):
    email: EmailStr
    password: str

class RegisterRequest(BaseModel):
    name: str
    email: EmailStr
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: dict

# === Model pentru cererea de chat ===
class ChatRequest(BaseModel):
    message: str
    page_context: dict = None  # Context opțional despre pagina unde se află chatul
    pdf_text: str = None  # Text extras din PDF
    session_id: Optional[int] = None  # ID-ul sesiunii de chat (opțional)
    user_id: Optional[int] = None  # ID-ul utilizatorului (opțional)

# === Funcții pentru autentificare ===
def hash_password(password: str) -> str:
    """Hash-uiește o parolă folosind bcrypt"""
    salt = bcrypt.gensalt()
    hashed = bcrypt.hashpw(password.encode('utf-8'), salt)
    return hashed.decode('utf-8')

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verifică dacă parola este corectă"""
    try:
        return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))
    except Exception as e:
        print(f"❌ Eroare la verificarea parolei: {e}")
        return False

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    """Creează un JWT token"""
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(hours=JWT_EXPIRATION_HOURS)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM)
    return encoded_jwt

def verify_token(token: str) -> Optional[dict]:
    """Verifică și decodează un JWT token"""
    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        print("❌ Token expirat")
        return None
    except jwt.InvalidTokenError:
        print("❌ Token invalid")
        return None

async def get_current_user(credentials: Optional[HTTPAuthorizationCredentials] = Depends(security_optional)) -> Optional[dict]:
    """Obține utilizatorul curent din token (opțional - permite și request-uri fără autentificare)"""
    if credentials is None:
        return None
    token = credentials.credentials
    payload = verify_token(token)
    if payload is None:
        return None
    user_id = payload.get("sub")
    if user_id is None:
        return None
    user = get_user(user_id=int(user_id))
    if user is None:
        return None
    # Elimină parola din răspuns
    if 'password' in user:
        del user['password']
    return user

# === Cache pentru config-uri ===
def get_cached_config(chat_id: str):
    """Obține config-ul din cache sau din baza de date"""
    # Verifică cache-ul
    if chat_id in _config_cache:
        return _config_cache[chat_id]
    
    # Încarcă din baza de date
    db_config = get_client_chat(chat_id)
    if not db_config:
        return None
    
    # Încarcă conținutul RAG din baza de date
    client_chat_id = db_config.get("id")
    rag_files_db = get_rag_files(client_chat_id, include_content=True)
    
    # Construiește rag_content din baza de date
    rag_content = []
    for rf in rag_files_db:
        if rf.get("content"):
            rag_content.append({
                "filename": rf.get("file", ""),
                "content": rf.get("content", "")
            })
    
    # Convertește la formatul așteptat
    config = {
        "name": db_config.get("name", "Chat nou"),
        "tenant_id": str(db_config.get("id", chat_id)),
        "model": db_config.get("model", "gpt-oss:20b"),
        "prompt": db_config.get("prompt", ""),
        "chat_title": db_config.get("chat_title"),
        "chat_subtitle": db_config.get("chat_subtitle"),
        "chat_color": db_config.get("chat_color", "#3b82f6"),
        "rag_files": db_config.get("rag_files", []),
        "rag_content": rag_content,  # Conținutul RAG din baza de date
        "institution": db_config.get("institution"),
        "created_at": db_config.get("created_at"),
        "updated_at": db_config.get("updated_at"),
        "is_active": bool(db_config.get("is_active", True))
    }
    
    # Salvează în cache
    _config_cache[chat_id] = config
    
    return config

# === Construiește prompt optimizat pentru JSON (o singură dată) ===
def build_json_instructions():
    """Construiește instrucțiunile JSON o singură dată (cache)"""
    return """JSON RAPID: Răspunde DOAR cu JSON valid, fără text. Chei: normalizează numele (lowercase, fără diacritice, spații→_). SELECT: folosește doar valori din opțiuni. OBLIGATORIU (*): completează întotdeauna. Format: {"cheie":"valoare"} - doar JSON pur."""

# Cache pentru instrucțiuni JSON
_JSON_INSTRUCTIONS = build_json_instructions()

# === Îmbunătățire prompt pentru detecție automată (optimizat) ===
def enhance_prompt_for_autofill(base_prompt, page_context=None, pdf_text=None, rag_content=None, institution_data=None, rag_search_query=None, tenant_id=None):
    """
    Îmbunătățește prompt-ul bazat pe contextul paginii, textul din PDF, conținutul RAG și datele instituției
    OPTIMIZAT: Folosește cache și format compact
    """
    # Dacă avem tenant_id și query pentru RAG, folosește vector store
    rag_context_text = None
    if tenant_id and rag_search_query:
        try:
            rag_store = get_tenant_rag_store(tenant_id)
            rag_results = rag_store.search(rag_search_query, top_k=5)
            if rag_results:
                rag_context_parts = []
                for result in rag_results:
                    rag_context_parts.append(f"\n--- {result['filename']} ---\n{result['content'][:2000]}")  # Limitează la 2000 caractere per chunk
                rag_context_text = "\n".join(rag_context_parts)
                print(f"✅ RAG search pentru tenant {tenant_id}: {len(rag_results)} rezultate relevante")
        except Exception as e:
            print(f"⚠️ Eroare la căutarea RAG pentru tenant {tenant_id}: {e}")
    
    # Dacă nu am folosit vector store, folosește rag_content direct
    if not rag_context_text and rag_content:
        rag_text = ""
        total_chars = 0
        max_total = 15000  # Mărim limita totală pentru RAG
        
        for item in rag_content:
            filename = item.get("filename", "document")
            content = item.get("content", "").strip()
            
            # Skip dacă conținutul este gol sau doar whitespace
            if not content or content == "\n":
                continue
            
            # Calculează cât mai putem adăuga
            remaining = max_total - total_chars
            if remaining <= 0:
                break
            
            # Limitează conținutul per fișier dar păstrăm mai mult (5000 per fișier)
            content_limited = content[:5000] if len(content) > 5000 else content
            
            # Verifică dacă mai avem spațiu
            if total_chars + len(content_limited) + len(filename) + 50 > max_total:
                # Adaugă doar cât mai încape
                available = max_total - total_chars - len(filename) - 50
                if available > 100:  # Doar dacă mai avem cel puțin 100 caractere
                    content_limited = content[:available]
                else:
                    break
            
            rag_text += f"\n\n--- {filename} ---\n{content_limited}"
            total_chars += len(content_limited) + len(filename) + 50
        
        if rag_text:
            rag_context_text = rag_text
            print(f"✅ RAG content adăugat în prompt: {len(rag_text)} caractere din {len(rag_content)} fișiere")
        else:
            print(f"⚠️ RAG content este gol sau invalid. Fișiere procesate: {len(rag_content) if rag_content else 0}")
    
    # Folosește prompt builder pentru generarea dinamică
    enhanced = build_dynamic_system_prompt(
        base_prompt=base_prompt,
        institution_data=institution_data,
        rag_context=rag_context_text
    )
    
    # Adaugă textul din PDF/imagini dacă există (format compact pentru viteză)
    if pdf_text:
        # Limitează la primele 2000 caractere pentru prompt (optimizare viteză mai agresivă)
        pdf_text_limited = pdf_text[:2000] if len(pdf_text) > 2000 else pdf_text
        enhanced += f"\n\n=== DOCUMENT ÎNCĂRCAT DE UTILIZATOR ===\n{pdf_text_limited}\n\nExtrage: nume, adrese, date, numere. Completează câmpurile pe baza acestui document."
    
    if page_context and page_context.get("has_form"):
        # Folosește informațiile detaliate despre câmpuri dacă sunt disponibile
        fields_detailed = page_context.get("fields_detailed", [])
        
        if fields_detailed:
            # Construiește descriere foarte compactă a câmpurilor (optimizat maxim pentru viteză)
            fields_list = []
            for field in fields_detailed[:30]:  # Limitează la primele 30 câmpuri
                field_info = field['name']
                if field.get('options'):
                    # Limitează la primele 2 opțiuni pentru viteză maximă
                    opts = ', '.join(field['options'][:2])
                    if len(field.get('options', [])) > 2:
                        opts += "..."
                    field_info += f" [{opts}]"
                if field.get('required'):
                    field_info += " *"
                fields_list.append(field_info)
            
            # Limitează lungimea totală a string-ului pentru prompt
            fields_str = ", ".join(fields_list)
            if len(fields_str) > 1500:  # Limitează la 1500 caractere
                fields_str = fields_str[:1500] + "..."
            
            # Folosește instrucțiunile din cache
            enhanced += f"\n\n=== CÂMPURI FORMULAR ===\n{fields_str}\n\n{_JSON_INSTRUCTIONS}"
        else:
            # Fallback la versiunea simplă dacă nu avem detalii (optimizat)
            fields_info = ", ".join(page_context.get("form_fields", [])[:30])  # Limitează la 30
            if len(fields_info) > 1000:
                fields_info = fields_info[:1000] + "..."
            enhanced += f"\n\n=== CÂMPURI FORMULAR ===\n{fields_info}\n\n{_JSON_INSTRUCTIONS}"
    
    return enhanced

# === Funcții helper pentru gestionarea contextului ===
def estimate_tokens(text: str) -> int:
    """Estimează numărul de tokens (aproximativ: ~4 caractere = 1 token)"""
    return len(text) // 4

def trim_conversation_history(history: list, max_chars: int = MAX_CONTEXT_CHARS - CONTEXT_RESERVE) -> list:
    """
    Taie istoricul conversației din început dacă depășește limita.
    Păstrează întotdeauna ultimele mesaje (user + assistant).
    """
    if not history:
        return history
    
    # Calculează lungimea totală
    total_chars = sum(len(msg.get("content", "")) for msg in history)
    
    if total_chars <= max_chars:
        return history
    
    # Taie din început, dar păstrează cel puțin ultimele 2 mesaje (user + assistant)
    trimmed = []
    current_chars = 0
    
    # Pornește de la sfârșit și adaugă mesaje până când depășim limita
    for msg in reversed(history):
        msg_chars = len(msg.get("content", ""))
        if current_chars + msg_chars > max_chars and len(trimmed) >= 2:
            # Am adăugat deja cel puțin 2 mesaje, oprește-te
            break
        trimmed.insert(0, msg)
        current_chars += msg_chars
    
    # Dacă tot e prea lung, taie din începutul listei trunchiate
    while current_chars > max_chars and len(trimmed) > 2:
        removed = trimmed.pop(0)
        current_chars -= len(removed.get("content", ""))
    
    print(f"✂️ Istoric trunchiat: {len(history)} -> {len(trimmed)} mesaje ({current_chars} caractere)")
    return trimmed

def get_conversation_history(chat_id: str = None, session_id: int = None, user_id: int = None) -> list:
    """Obține istoricul conversației pentru un chat_id sau session_id din baza de date"""
    # Dacă avem session_id, folosim sesiunea (mod nou)
    if session_id:
        messages = db_get_conversation_history(chat_id=None, session_id=session_id, user_id=user_id)
        return messages
    
    # Dacă avem chat_id, folosim modul vechi (compatibilitate)
    if chat_id:
        # Încearcă să folosească cache-ul în memorie pentru performanță (opțional)
        tenant_id = get_tenant_id_from_chat_id(chat_id)
        if tenant_id in _conversation_history and chat_id in _conversation_history[tenant_id]:
            return _conversation_history[tenant_id][chat_id]
        
        # Încarcă din baza de date
        messages = db_get_conversation_history(chat_id=chat_id, session_id=None, user_id=user_id)
        
        # Salvează în cache pentru performanță
        if tenant_id not in _conversation_history:
            _conversation_history[tenant_id] = {}
        _conversation_history[tenant_id][chat_id] = messages
        
        return messages
    
    return []

def add_to_conversation_history(chat_id: str = None, session_id: int = None, role: str = None, content: str = None, user_id: int = None):
    """Adaugă un mesaj la istoricul conversației (salvează în DB și cache)"""
    # Salvează în baza de date
    db_add_message_to_conversation(session_id=session_id, chat_id=chat_id, role=role, content=content, user_id=user_id)
    
    # Actualizează cache-ul doar pentru modul vechi (chat_id fără session_id)
    if chat_id and not session_id:
        tenant_id = get_tenant_id_from_chat_id(chat_id)
        if tenant_id not in _conversation_history:
            _conversation_history[tenant_id] = {}
        if chat_id not in _conversation_history[tenant_id]:
            _conversation_history[tenant_id][chat_id] = []
    
        _conversation_history[tenant_id][chat_id].append({
        "role": role,
        "content": content
    })
    
    # Aplică limitarea contextului
        _conversation_history[tenant_id][chat_id] = trim_conversation_history(_conversation_history[tenant_id][chat_id])

def clear_conversation_history(chat_id: str, user_id: int = None):
    """Șterge istoricul conversației pentru un chat_id (din DB și cache)"""
    # Șterge din baza de date
    db_clear_conversation_history(chat_id, user_id)
    
    # Șterge din cache
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    if tenant_id in _conversation_history and chat_id in _conversation_history[tenant_id]:
        del _conversation_history[tenant_id][chat_id]
        print(f"🗑️ Istoric șters pentru chat_id: {chat_id} (tenant: {tenant_id})")

def create_default_config(chat_id: str):
    """Creează un config default pentru un chat_id dacă nu există"""
    # Verifică dacă deja există
    existing = get_cached_config(chat_id)
    if existing:
        return existing
    
    # Creează în baza de date
    chat_id_int = create_client_chat(
        name="Chat nou",
        model="gpt-oss:20b",
        prompt="Ești asistentul Integra AI. Răspunde clar și politicos la întrebările utilizatorilor.",
        chat_title="Chat nou",
        chat_subtitle="Asistentul tău inteligent pentru găsirea informațiilor",
        chat_color="#3b82f6"
    )
    
    if not chat_id_int:
        print(f"❌ Nu s-a putut crea config pentru {chat_id}")
        return None
    
    # Reîncarcă config-ul creat
    config = get_cached_config(str(chat_id_int))
    
    print(f"✅ Config default creat pentru chat_id: {chat_id_int}")
    return config

# === Stream răspuns cu prompt îmbunătățit ===
async def stream_response(messages, model, page_context=None, pdf_text=None, rag_content=None, institution_data=None, rag_search_query=None, tenant_id=None):
    # Îmbunătățește primul mesaj (system prompt) dacă există context
    if len(messages) > 0:
        messages[0]['content'] = enhance_prompt_for_autofill(
            messages[0]['content'], 
            page_context,
            pdf_text,
            rag_content,
            institution_data,
            rag_search_query,
            tenant_id
        )
    
    # Parametrii optimizați pentru viteză (doar când există context de formular)
    options = {}
    if page_context and page_context.get("has_form"):
        # Optimizări doar pentru generare JSON (mai rapid)
        options = {
            "temperature": 0.2,  # Foarte determinist pentru JSON rapid
            "top_p": 0.85,  # Redus pentru generare mai rapidă
            "top_k": 20,  # Limitează opțiunile pentru viteză
            "num_predict": 2000,  # Suficient pentru JSON complet
            "repeat_penalty": 1.1,  # Evită repetări
        }
    
    stream = ollama.chat(
        model=model, 
        messages=messages, 
        stream=True,
        options=options if options else None
    )
    
    for chunk in stream:
        if "message" in chunk and "content" in chunk["message"]:
            content = chunk["message"]["content"]
            if content:
                yield content
        await asyncio.sleep(0)


# === Endpoint pentru chaturi dinamice ===
@app.post("/chat/{chat_id}/ask")
async def ask_dynamic(chat_id: str, request: ChatRequest, current_user: dict = Depends(get_current_user)):
    # Verifică autentificarea (opțional - poate fi dezactivat pentru guest users)
    # Dacă nu există current_user, folosește user_id din request sau default
    user_id = None
    if current_user:
        user_id = current_user.get('id')
    else:
        user_id = getattr(request, 'user_id', None) or 1  # Default user_id = 1 pentru guest
    
    # Folosește cache pentru config (mult mai rapid)
    config = get_cached_config(chat_id)
    
    # Dacă config-ul nu există, creează unul default
    if not config:
        print(f"⚠️ Config nu există pentru {chat_id}, creez config default...")
        config = create_default_config(chat_id)
    
    # Extrage session_id din request (dacă există)
    session_id = getattr(request, 'session_id', None)
    
    # Dacă avem session_id, folosim sesiunea; altfel folosim modul vechi (compatibilitate)
    if session_id:
        # Verifică dacă sesiunea există
        session = get_chat_session(session_id)
        if not session:
            return JSONResponse(
                status_code=404,
                content={"error": f"Session {session_id} not found"}
            )
        # Verifică dacă sesiunea aparține user-ului curent (dacă este autentificat)
        if current_user and session.get('user_id') != current_user.get('id'):
            return JSONResponse(
                status_code=403,
                content={"error": "Nu ai acces la această sesiune de chat"}
            )
        # Folosește user_id din sesiune
        user_id = session.get('user_id', user_id)
    else:
        # Dacă nu avem session_id dar avem user_id, creează o sesiune nouă
        try:
            client_chat_id = int(chat_id)
        except ValueError:
            # Dacă nu este int, caută după name
            from database import get_client_chat
            client = get_client_chat(chat_id)
            if not client:
                return JSONResponse(
                    status_code=404,
                    content={"error": f"Chat {chat_id} not found"}
                )
            client_chat_id = client['id']
        
        # Creează o sesiune nouă (funcția create_chat_session va crea automat user-ul dacă nu există)
        session_id = create_chat_session(user_id, client_chat_id, None)
        if not session_id:
            return JSONResponse(
                status_code=500,
                content={"error": "Failed to create chat session"}
            )
        print(f"✅ Sesiune nouă creată: {session_id} pentru user {user_id}, chat {chat_id}")
    
    # Extrage conținutul RAG din config dacă există
    rag_content = config.get("rag_content", [])
    rag_files = config.get("rag_files", [])
    
    # Log pentru debugging
    if rag_content:
        valid_rag = [r for r in rag_content if r.get("content", "").strip() and r.get("content", "") != "\n"]
        print(f"📚 RAG pentru {chat_id}: {len(valid_rag)} fișiere valide din {len(rag_content)} totale")
        if len(valid_rag) == 0:
            if rag_files:
                print(f"⚠️ ATENȚIE: Există {len(rag_files)} fișiere RAG pentru {chat_id}, dar toate sunt goale (probabil PDF-uri scanate). Re-procesează cu OCR sau convertește manual la text.")
            else:
                print(f"⚠️ ATENȚIE: Toate fișierele RAG sunt goale pentru {chat_id}! Re-procesează fișierele.")
    else:
        if rag_files:
            print(f"ℹ️ Există {len(rag_files)} fișiere RAG pentru {chat_id}, dar nu au conținut extractibil (probabil PDF-uri scanate). Re-procesează cu OCR.")
        else:
            print(f"ℹ️ Nu există RAG content pentru {chat_id}")
    
    # === GESTIONARE ISTORIC CONVERSAȚIE ===
    # Obține istoricul existent folosind session_id sau chat_id (compatibilitate)
    conversation_history = db_get_conversation_history(chat_id=chat_id if not session_id else None, session_id=session_id, user_id=user_id)
    
    # Adaugă mesajul nou al utilizatorului în istoric
    user_message = request.message
    db_add_message_to_conversation(session_id=session_id, chat_id=chat_id if not session_id else None, role="user", content=user_message, user_id=user_id)
    
    # Obține istoricul actualizat
    updated_history = db_get_conversation_history(chat_id=chat_id if not session_id else None, session_id=session_id, user_id=user_id)
    
    # Extrage datele instituției și tenant_id
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    institution_data = config.get("institution")
    
    # Construiește mesajele cu istoricul complet
    # System prompt-ul va fi generat dinamic în stream_response
    messages = [{"role": "system", "content": config["prompt"]}]
    
    # Adaugă istoricul conversației complet (inclusiv mesajul nou al utilizatorului)
    messages.extend(updated_history)
    
    # Log pentru debugging
    print(f"💬 Conversație pentru {chat_id} (session: {session_id}, tenant: {tenant_id}): {len(conversation_history)} mesaje istorice + 1 mesaj nou = {len(updated_history)} mesaje totale în context")
    
    # === STREAM RĂSPUNS CU COLECTARE ===
    # Folosim un wrapper care colectează răspunsul complet
    full_response = ""
    
    # Folosește mesajul utilizatorului pentru căutare RAG semantică
    rag_search_query = request.message if request.message else None
    
    async def stream_with_collection():
        nonlocal full_response
        async for chunk in stream_response(
            messages, 
            config["model"], 
            request.page_context, 
            request.pdf_text, 
            rag_content,
            institution_data,
            rag_search_query,
            tenant_id
        ):
            full_response += chunk
            yield chunk
        
        # După ce s-a terminat streaming-ul, salvează răspunsul în istoric
        if full_response.strip():
            db_add_message_to_conversation(session_id=session_id, chat_id=chat_id if not session_id else None, role="assistant", content=full_response, user_id=user_id)
            print(f"✅ Răspuns salvat în istoric pentru {chat_id} (session: {session_id}): {len(full_response)} caractere")
   
    return StreamingResponse(
        stream_with_collection(), 
        media_type="text/plain; charset=utf-8"
    )

# === Endpoint pentru obținere configurație chat ===
@app.get("/chat/{chat_id}/config")
async def get_chat_config(chat_id: str, current_user: dict = Depends(get_current_user)):
    # Verifică dacă chat-ul există
    # Folosește cache pentru config (mult mai rapid)
    config = get_cached_config(chat_id)
    
    # Dacă config-ul nu există, returnează 404 (nu creează automat)
    if not config:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Chat {chat_id} nu există"
        )
    
    # Nu returnăm conținutul RAG complet (prea mare), doar metadata
    response_config = {**config}
    if "rag_content" in response_config:
        # Returnăm doar informații despre fișiere, nu conținutul complet
        rag_info = []
        for item in response_config["rag_content"]:
            content = item.get("content", "")
            rag_info.append({
                "filename": item.get("filename", ""),
                "content_length": len(content),
                "has_content": bool(content and content.strip() and content != "\n")
            })
        response_config["rag_content_info"] = rag_info
        # Nu trimitem conținutul complet în response (prea mare pentru frontend)
        del response_config["rag_content"]
    
    return JSONResponse(content=response_config)

# === Endpoint pentru creare sesiune de chat ===
@app.post("/chat/{chat_id}/session/create")
async def create_session(chat_id: str, request: dict, current_user: dict = Depends(get_current_user)):
    """Creează o nouă sesiune de chat pentru un utilizator"""
    try:
        # Verifică dacă chat-ul există
        config = get_cached_config(chat_id)
        if not config:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Chat {chat_id} nu există"
            )
        
        user_id = current_user.get('id') if current_user else request.get("user_id", 1)
        title = request.get("title", None)
        
        # Convertește chat_id la int
        try:
            client_chat_id = int(chat_id)
        except ValueError:
            # Dacă nu este int, caută după name
            client = get_client_chat(chat_id)
            if not client:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"Chat {chat_id} nu există"
                )
            client_chat_id = client['id']
        
        session_id = create_chat_session(user_id, client_chat_id, title)
        if not session_id:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Eroare la crearea sesiunii de chat"
            )
        
        return JSONResponse(content={
            "success": True,
            "session_id": session_id,
            "message": f"Sesiune de chat creată cu succes"
        })
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Eroare la crearea sesiunii: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )

# === Endpoint pentru listare sesiuni de chat ===
@app.get("/chat/{chat_id}/sessions")
async def list_sessions(chat_id: str, current_user: dict = Depends(get_current_user)):
    """Listează toate sesiunile de chat ale unui utilizator pentru un chatbot"""
    try:
        # Verifică dacă chat-ul există
        config = get_cached_config(chat_id)
        if not config:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Chat {chat_id} nu există"
            )
        
        # Convertește chat_id la int
        try:
            client_chat_id = int(chat_id)
        except ValueError:
            # Dacă nu este int, caută după name
            client = get_client_chat(chat_id)
            if not client:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"Chat {chat_id} nu există"
                )
            client_chat_id = client['id']
        
        # Obține user_id din token sau default
        user_id = current_user.get('id') if current_user else 1
        
        sessions = list_user_chat_sessions(user_id, client_chat_id)
        
        return JSONResponse(content={
            "success": True,
            "sessions": sessions
        })
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Eroare la listarea sesiunilor: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )

# === Endpoint pentru resetare conversație ===
@app.post("/chat/{chat_id}/clear")
async def clear_chat_history(chat_id: str, session_id: Optional[int] = None, current_user: dict = Depends(get_current_user)):
    """Șterge istoricul conversației pentru un chat sau o sesiune"""
    # Dacă avem session_id, verifică că sesiunea aparține user-ului
    if session_id:
        session = get_chat_session(session_id)
        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Sesiune {session_id} nu există"
            )
        if current_user and session.get('user_id') != current_user.get('id'):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Nu ai acces la această sesiune"
            )
    
    db_clear_conversation_history(session_id=session_id, chat_id=chat_id if not session_id else None)
    return JSONResponse(content={
        "success": True,
        "message": f"Istoricul conversației a fost șters"
    })

# === Endpoint pentru actualizare sesiune (redenumire) ===
@app.put("/chat/{chat_id}/session/{session_id}")
async def update_session(chat_id: str, session_id: int, request: dict, current_user: dict = Depends(get_current_user)):
    """Actualizează o sesiune de chat (redenumire)"""
    try:
        # Verifică că sesiunea există
        session = get_chat_session(session_id)
        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Sesiune {session_id} nu există"
            )
        
        # Verifică accesul
        if current_user and session.get('user_id') != current_user.get('id'):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Nu ai acces la această sesiune"
            )
        
        # Extrage title din request
        title = request.get('title')
        if title is None:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Câmpul 'title' este obligatoriu"
            )
        
        # Actualizează sesiunea
        success = db_update_chat_session(session_id, title)
        if not success:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Eroare la actualizarea sesiunii"
            )
        
        return JSONResponse(content={
            "success": True,
            "message": "Sesiunea a fost actualizată cu succes",
            "session": get_chat_session(session_id)
        })
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Eroare la actualizarea sesiunii: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )

# === Endpoint pentru ștergere sesiune ===
@app.delete("/chat/{chat_id}/session/{session_id}")
async def delete_session(chat_id: str, session_id: int, current_user: dict = Depends(get_current_user)):
    """Șterge o sesiune de chat și toate mesajele asociate"""
    try:
        # Verifică că sesiunea există
        session = get_chat_session(session_id)
        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Sesiune {session_id} nu există"
            )
        
        # Verifică accesul
        if current_user and session.get('user_id') != current_user.get('id'):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Nu ai acces la această sesiune"
            )
        
        # Șterge sesiunea (mesajele se șterg automat prin CASCADE)
        success = db_delete_chat_session(session_id)
        if not success:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Eroare la ștergerea sesiunii"
            )
        
        return JSONResponse(content={
            "success": True,
            "message": "Sesiunea și toate mesajele au fost șterse cu succes"
        })
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Eroare la ștergerea sesiunii: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )

# === Endpoint pentru obținere istoric conversație ===
@app.get("/chat/{chat_id}/history")
async def get_chat_history(chat_id: str, session_id: Optional[int] = None, current_user: dict = Depends(get_current_user)):
    """Obține istoricul conversației pentru frontend"""
    # Extrage user_id
    user_id = None
    if current_user:
        user_id = current_user.get('id')
    else:
        # Dacă nu este autentificat, folosește default user
        user_id = 1  # Default user
    
    # Dacă avem session_id, verifică că sesiunea aparține user-ului
    if session_id:
        session = get_chat_session(session_id)
        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Sesiune {session_id} nu există"
            )
        # Verifică accesul (dacă este autentificat)
        if current_user and session.get('user_id') != current_user.get('id'):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Nu ai acces la această sesiune"
            )
        # Folosește user_id din sesiune
        user_id = session.get('user_id', user_id)
    
    # Obține istoricul
    history = db_get_conversation_history(chat_id=chat_id if not session_id else None, session_id=session_id, user_id=user_id)
    
    return JSONResponse(content={
        "chat_id": chat_id,
        "session_id": session_id,
        "message_count": len(history),
        "messages": history
    })

# === Endpoint pentru re-procesare fișiere RAG ===
@app.post("/chat/{chat_id}/reprocess-rag")
async def reprocess_rag(chat_id: str):
    """Re-procesează fișierele RAG pentru un chat existent"""
    config = get_cached_config(chat_id)
    if not config:
        return JSONResponse(
            status_code=404,
            content={"error": f"Chat configuration not found: {chat_id}"}
        )
    
    rag_dir = f"rag/{chat_id}"
    
    if not os.path.exists(rag_dir):
        return JSONResponse(
            status_code=404,
            content={"error": f"Directorul RAG nu există: {rag_dir}"}
        )
    
    # Dacă rag_files este gol, detectează automat fișierele din director
    rag_files = config.get("rag_files", [])
    if not rag_files:
        # Detectează automat toate fișierele din director
        if os.path.exists(rag_dir):
            rag_files = [f for f in os.listdir(rag_dir) if os.path.isfile(os.path.join(rag_dir, f))]
            # Actualizează config-ul cu fișierele detectate (doar în cache, nu în DB)
            config["rag_files"] = rag_files
            print(f"✅ Detectate automat {len(rag_files)} fișiere RAG în {rag_dir}")
    
    if not rag_files:
        return JSONResponse(
            status_code=400,
            content={"error": "Nu există fișiere RAG pentru acest chat"}
        )
    
    rag_content = []
    
    for filename in rag_files:
        file_path = os.path.join(rag_dir, filename)
        if not os.path.exists(file_path):
            print(f"⚠️ Fișier RAG nu există: {file_path}")
            continue
        
        text_content = ""
        try:
            if filename.endswith('.pdf') and PDF_AVAILABLE:
                with open(file_path, "rb") as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content += f"\n--- Pagina {page_num + 1} ---\n{page_text}\n"
                        except Exception as e:
                            print(f"Eroare la extragerea paginii {page_num + 1} din {filename}: {e}")
                            continue
                
                # Dacă nu s-a extras text (PDF scanat), logăm
                if not text_content.strip() and OCR_AVAILABLE:
                    print(f"⚠️ PDF {filename} pare scanat sau nu conține text extractibil.")
            elif filename.endswith(('.txt', '.md')):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text_content = f.read()
                except UnicodeDecodeError:
                    try:
                        with open(file_path, "r", encoding="latin-1") as f:
                            text_content = f.read()
                    except Exception as e:
                        print(f"Eroare la citirea fișierului {filename}: {e}")
                except Exception as e:
                    print(f"Eroare la citirea fișierului {filename}: {e}")
            elif filename.endswith(('.doc', '.docx')):
                try:
                    try:
                        from docx import Document
                        doc = Document(file_path)
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text_content += para.text + "\n"
                        # Extrage și din tabele
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                if row_text.strip():
                                    text_content += row_text + "\n"
                    except ImportError:
                        print(f"⚠️ python-docx nu este instalat. Pentru DOC/DOCX, rulează: pip install python-docx")
                        text_content = f"[Fișier DOC/DOCX - instalează python-docx pentru extragere: pip install python-docx]"
                except Exception as e:
                    print(f"Eroare la extragerea textului din DOC/DOCX {filename}: {e}")
        except Exception as e:
            print(f"Eroare la procesarea {filename}: {e}")
        
        if text_content and text_content.strip():
            rag_content.append({
                "filename": filename,
                "content": text_content.strip()
            })
            print(f"✅ Text re-extras din {filename}: {len(text_content)} caractere")
    
    # Actualizează vector store-ul pentru tenant
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        # Șterge toate documentele vechi
        rag_store.clear()
        # Adaugă documentele noi
        for item in rag_content:
            rag_store.add_document(item["filename"], item["content"])
        print(f"✅ Vector store actualizat pentru tenant {tenant_id}")
    except Exception as e:
        print(f"⚠️ Eroare la actualizarea vector store pentru tenant {tenant_id}: {e}")
    
    # Invalidează cache-ul
    if chat_id in _config_cache:
        del _config_cache[chat_id]
    
    return JSONResponse(content={
        "success": True,
        "processed_files": len(rag_content),
        "total_files": len(rag_files),
        "message": f"Re-procesat {len(rag_content)} din {len(rag_files)} fișiere și actualizat vector store"
    })

# === Endpoint pentru listare chaturi disponibile ===
@app.get("/chats/list")
async def list_chats():
    # Încarcă din baza de date
    db_chats = list_all_client_chats()
    
    chats = []
    for db_chat in db_chats:
                chats.append({
            "id": str(db_chat.get("id", "")),
            "name": db_chat.get("name", "Unknown"),
            "model": db_chat.get("model", "unknown")
                })
    
    return JSONResponse(content={"chats": chats})
    
# === Pagina de chat dinamic ===
@app.get("/chat/{chat_id}", response_class=HTMLResponse)
async def serve_chat(chat_id: str):
    # Verifică dacă chat-ul există în baza de date
    config = get_cached_config(chat_id)
    if not config:
        return HTMLResponse("<h3>Chat configurat inexistent.</h3>")
    with open("public/index.html", "r", encoding="utf-8") as f:
        return HTMLResponse(f.read())

# === Formular pentru configurarea unui chat ===
@app.get("/builder", response_class=HTMLResponse)
async def builder_page():
    return FileResponse("public/builder.html")

# === Salvare configurare ===
@app.post("/builder/create")
async def create_chat(
    request: Request,
    name: str = Form(...),
    model: str = Form(...),
    prompt: str = Form(...),
    chat_title: Optional[str] = Form(None),
    chat_subtitle: Optional[str] = Form(None),
    chat_color: Optional[str] = Form("#3b82f6"),
    rag_files: Optional[list[UploadFile]] = File(None)
):
    # Generează ID unic
    chat_id = name.lower().replace(" ", "-") + "-" + str(uuid.uuid4())[:8]
    
    # Procesează fișierele RAG
    rag_content = []
    if rag_files:
        os.makedirs("rag", exist_ok=True)
        os.makedirs(f"rag/{chat_id}", exist_ok=True)
        for file in rag_files:
            if file.filename:
                # Salvează fișierul
                file_path = f"rag/{chat_id}/{file.filename}"
                with open(file_path, "wb") as f:
                    content = await file.read()
                    f.write(content)
                
                # Extrage text din fișier
                text_content = ""
                if file.filename.endswith('.pdf') and PDF_AVAILABLE:
                    try:
                        with open(file_path, "rb") as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            for page_num, page in enumerate(pdf_reader.pages):
                                try:
                                    page_text = page.extract_text()
                                    if page_text.strip():
                                        text_content += f"\n--- Pagina {page_num + 1} ---\n{page_text}\n"
                                except Exception as e:
                                    print(f"Eroare la extragerea paginii {page_num + 1} din {file.filename}: {e}")
                                    continue
                        
                        # Dacă nu s-a extras text (PDF scanat), încearcă OCR
                        if not text_content.strip() and OCR_AVAILABLE:
                            try:
                                print(f"PDF {file.filename} pare scanat, încerc OCR...")
                                # Pentru PDF-uri scanate, ar trebui să convertim paginile în imagini
                                # Pentru moment, doar logăm
                                print(f"⚠️ PDF {file.filename} nu conține text extractibil. Poate fi scanat.")
                            except Exception as e:
                                print(f"Eroare OCR pentru {file.filename}: {e}")
                    except Exception as e:
                        print(f"Eroare la extragerea textului din PDF {file.filename}: {e}")
                elif file.filename.endswith(('.txt', '.md')):
                    try:
                        with open(file_path, "r", encoding="utf-8") as f:
                            text_content = f.read()
                    except UnicodeDecodeError:
                        # Încearcă cu encoding diferit
                        try:
                            with open(file_path, "r", encoding="latin-1") as f:
                                text_content = f.read()
                        except Exception as e:
                            print(f"Eroare la citirea fișierului {file.filename}: {e}")
                    except Exception as e:
                        print(f"Eroare la citirea fișierului {file.filename}: {e}")
                elif file.filename.endswith(('.doc', '.docx')):
                    # Pentru DOC/DOCX, ar trebui python-docx
                    try:
                        try:
                            from docx import Document
                            doc = Document(file_path)
                            for para in doc.paragraphs:
                                if para.text.strip():
                                    text_content += para.text + "\n"
                            # Extrage și din tabele
                            for table in doc.tables:
                                for row in table.rows:
                                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                    if row_text.strip():
                                        text_content += row_text + "\n"
                        except ImportError:
                            print(f"⚠️ python-docx nu este instalat. Pentru DOC/DOCX, rulează: pip install python-docx")
                            text_content = f"[Fișier DOC/DOCX - instalează python-docx pentru extragere: pip install python-docx]"
                    except Exception as e:
                        print(f"Eroare la extragerea textului din DOC/DOCX {file.filename}: {e}")
                
                if text_content and text_content.strip():
                    # Nu limităm aici - limităm doar în prompt pentru a păstra tot conținutul în config
                    rag_content.append({
                        "filename": file.filename,
                        "content": text_content.strip()  # Păstrăm tot conținutul în config
                    })
                    print(f"✅ Text extras din {file.filename}: {len(text_content)} caractere")
                else:
                    print(f"⚠️ Nu s-a putut extrage text din {file.filename} (poate fi gol, scanat sau protejat)")
    
    # IMPORTANT: Nu modificăm prompt-ul salvat în config
    # Îmbunătățirea se face dinamic la runtime în funcție de context
    
    # Creează chatbot-ul în baza de date
    client_chat_id = create_client_chat(
        name=name,
        model=model,
        prompt=prompt,
        chat_title=chat_title or name,
        chat_subtitle=chat_subtitle or "Asistentul tău inteligent pentru găsirea informațiilor",
        chat_color=chat_color or "#3b82f6"
    )
    
    if not client_chat_id:
        return JSONResponse(
            status_code=500,
            content={"error": "Eroare la crearea chatbot-ului în baza de date"}
        )
    
    # Adaugă fișierele RAG în baza de date și vector store
    tenant_id = str(client_chat_id)
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        # Șterge toate documentele vechi
        rag_store.clear()
        
        # Adaugă documentele noi
        for item in rag_content:
            # Salvează în DB cu conținutul
            add_rag_file(client_chat_id, item["filename"], item["content"])
            # Adaugă în vector store
            rag_store.add_document(item["filename"], item["content"])
        
        print(f"✅ Vector store creat pentru tenant {tenant_id}")
    except Exception as e:
        print(f"⚠️ Eroare la crearea vector store pentru tenant {tenant_id}: {e}")
    
    # Reîncarcă config-ul din DB
    config = get_cached_config(str(client_chat_id))

    # Returnează link direct pentru chat full-page
    chat_url = f"/chat/{client_chat_id}"
    
    return JSONResponse(content={
        "chat_id": str(client_chat_id),
        "chat_url": chat_url,
        "config": config
    })

# === Endpoint default pentru test ===
@app.post("/ask")
async def ask_default(request: ChatRequest):
    default_config = {
        "model": "gpt-oss:20b",
        "prompt": "Ești asistentul Integra AI. Răspunde clar și politicos."
    }

    messages = [
        {"role": "system", "content": default_config["prompt"]},
        {"role": "user", "content": request.message}
    ]

    return StreamingResponse(
        stream_response(messages, default_config["model"], request.page_context, request.pdf_text, None, None, None, None),
        media_type="text/plain; charset=utf-8"
    )

@app.get("/")
async def serve_index():
    return FileResponse("public/index.html")

@app.get("/scripts/{filename}")
async def serve_script(filename: str):
    file_path = os.path.join("public", "scripts", filename)
    if not os.path.exists(file_path):
        print(f"⚠️ Fișier script lipsă: {file_path}")
        return JSONResponse(status_code=404, content={"error": "File not found"})
    return FileResponse(file_path)

@app.get("/style/{filename}")
async def serve_style(filename: str):
    file_path = os.path.join("public", "style", filename)
    if not os.path.exists(file_path):
        print(f"⚠️ Fișier CSS lipsă: {file_path}")
        return JSONResponse(status_code=404, content={"error": "File not found"})
    return FileResponse(file_path)

@app.get("/favicon.ico")
async def favicon():
    # Return empty response pentru favicon (browserul va încerca să-l încarce)
    return Response(status_code=204)

# === Pagini HTML statice ===
@app.get("/evenimente", response_class=HTMLResponse)
async def serve_events():
    return FileResponse("public/formular_evenimente.html")

@app.get("/primarie", response_class=HTMLResponse)
async def serve_primarie():
    return FileResponse("public/site_primarie.html")

@app.get("/rezervari", response_class=HTMLResponse)
async def serve_reservations():
    return FileResponse("public/formular_rezervari.html")

# === Ruta pentru pagina de login (Next.js) ===
@app.get("/login")
async def serve_login():
    """
    Redirecționează către pagina de login Next.js
    Dacă Next.js rulează pe același server, folosește redirect.
    Dacă rulează separat, ajustează URL-ul corespunzător.
    """
    # Construiește URL-ul pentru redirect din variabila de mediu
    nextjs_base_url = os.getenv('NEXTJS_URL', 'http://localhost:3000')
    nextjs_url = f"{nextjs_base_url}/login"
    return RedirectResponse(url=nextjs_url, status_code=307)

# === Ruta pentru pagina de înregistrare (Next.js) ===
@app.get("/register")
async def serve_register():
    """
    Redirecționează către pagina de înregistrare Next.js
    """
    nextjs_base_url = os.getenv('NEXTJS_URL', 'http://localhost:3000')
    nextjs_url = f"{nextjs_base_url}/register"
    return RedirectResponse(url=nextjs_url, status_code=307)

# === Endpoint pentru extragerea textului din PDF ===
@app.post("/extract-pdf")
async def extract_pdf(pdf: UploadFile = File(...)):
    """
    Extrage textul dintr-un fișier PDF
    """
    if not PDF_AVAILABLE:
        return JSONResponse(
            status_code=500,
            content={"error": "PyPDF2 nu este instalat. Rulează: pip install PyPDF2"}
        )
    
    if pdf.content_type != "application/pdf":
        return JSONResponse(
            status_code=400,
            content={"error": "Fișierul trebuie să fie PDF"}
        )
    
    try:
        # Citește conținutul PDF
        pdf_content = await pdf.read()
        
        # Extrage textul folosind PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                text += f"\n--- Pagina {page_num + 1} ---\n"
                text += page_text
            except Exception as e:
                print(f"Eroare la extragerea paginii {page_num + 1}: {e}")
                continue
        
        if not text.strip():
            return JSONResponse(
                status_code=400,
                content={"error": "Nu s-a putut extrage text din PDF. PDF-ul poate fi scanat sau protejat."}
            )
        
        return JSONResponse(content={
            "text": text.strip(),
            "pages": len(pdf_reader.pages),
            "filename": pdf.filename
        })
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la procesarea PDF: {str(e)}"}
        )

# === Endpoint pentru extragerea textului din imagini (OCR) ===
@app.post("/extract-image")
async def extract_image(image: UploadFile = File(...)):
    """
    Extrage textul dintr-o imagine folosind OCR
    """
    if not OCR_AVAILABLE:
        return JSONResponse(
            status_code=500,
            content={"error": "OCR nu este disponibil. Rulează: pip install pytesseract pillow. Asigură-te că Tesseract OCR este instalat pe sistem."}
        )
    
    # Verifică tipul de fișier
    allowed_types = ["image/jpeg", "image/jpg", "image/png", "image/gif", "image/bmp", "image/webp"]
    if image.content_type not in allowed_types:
        return JSONResponse(
            status_code=400,
            content={"error": f"Fișierul trebuie să fie o imagine. Tip primit: {image.content_type}. Tipuri acceptate: {', '.join(allowed_types)}"}
        )
    
    try:
        # Citește conținutul imaginii
        image_content = await image.read()
        
        if not image_content:
            return JSONResponse(
                status_code=400,
                content={"error": "Fișierul este gol sau nu a putut fi citit."}
            )
        
        # Deschide imaginea cu PIL
        try:
            img = Image.open(io.BytesIO(image_content))
        except Exception as e:
            return JSONResponse(
                status_code=400,
                content={"error": f"Nu s-a putut deschide imaginea: {str(e)}"}
            )
        
        # Extrage textul folosind OCR
        text = None
        error_msg = None
        
        # Verifică dacă Tesseract este disponibil încă o dată (în caz că OCR_AVAILABLE era True dar Tesseract nu funcționează)
        try:
            pytesseract.get_tesseract_version()
        except Exception as tess_check_error:
            error_msg = str(tess_check_error)
            if "tesseract" in error_msg.lower() or "not found" in error_msg.lower() or "no such file" in error_msg.lower():
                return JSONResponse(
                    status_code=500,
                    content={"error": f"Tesseract OCR nu este instalat sau nu este în PATH. Eroare: {error_msg}. Instalează Tesseract OCR de la: https://github.com/UB-Mannheim/tesseract/wiki"}
                )
        
        # Încearcă cu diferite configurații de limbi
        lang_configs = ['ron+eng', 'eng', 'ron', None]  # None = default
        
        for lang_config in lang_configs:
            try:
                if lang_config:
                    text = pytesseract.image_to_string(img, lang=lang_config)
                else:
                    text = pytesseract.image_to_string(img)
                
                if text and text.strip():
                    break  # Dacă am obținut text, ieșim din loop
            except Exception as e:
                error_msg = str(e)
                # Dacă e eroare de Tesseract, oprește imediat
                if "tesseract" in error_msg.lower() or "not found" in error_msg.lower() or "no such file" in error_msg.lower():
                    return JSONResponse(
                        status_code=500,
                        content={"error": f"Tesseract OCR nu este instalat sau nu este în PATH. Eroare: {error_msg}. Instalează Tesseract OCR de la: https://github.com/UB-Mannheim/tesseract/wiki"}
                    )
                # Continuă cu următoarea configurație pentru alte erori
                continue
        
        # Dacă nu am reușit să extragem text, verifică eroarea
        if not text or not text.strip():
            if error_msg:
                if "tesseract" in error_msg.lower() or "not found" in error_msg.lower() or "no such file" in error_msg.lower():
                    return JSONResponse(
                        status_code=500,
                        content={"error": f"Tesseract OCR nu este instalat sau nu este în PATH. Eroare: {error_msg}. Vezi INSTALARE_OCR.md pentru instrucțiuni."}
                    )
                else:
                    return JSONResponse(
                        status_code=500,
                        content={"error": f"Eroare la extragerea textului cu OCR: {error_msg}"}
                    )
            else:
                # Nu am eroare, dar nici text - probabil imaginea nu conține text
                return JSONResponse(
                    status_code=400,
                    content={"error": "Nu s-a putut extrage text din imagine. Imaginea poate să nu conțină text sau calitatea este prea slabă. Încearcă cu o imagine de calitate mai bună."}
                )
        
        return JSONResponse(content={
            "text": text.strip(),
            "filename": image.filename,
            "type": "image"
        })
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"❌ Eroare detaliată la procesarea imaginii: {error_details}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la procesarea imaginii: {str(e)}. Verifică consola serverului pentru detalii."}
        )

# === Endpoint-uri pentru administrarea tenant-ilor ===

@app.put("/admin/tenant/{chat_id}/institution")
async def update_institution_data(chat_id: str, institution_data: dict):
    """Actualizează datele instituției pentru un tenant"""
    # Convertește chat_id la int
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        # Caută după name
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Salvează în baza de date
    success = create_or_update_client_type(
        client_chat_id=client_chat_id,
        name=institution_data.get("name", ""),
        type=institution_data.get("type", "alta"),
        address=institution_data.get("address"),
        phone=institution_data.get("phone"),
        email=institution_data.get("email"),
        website=institution_data.get("website")
    )
    
    if not success:
        return JSONResponse(
            status_code=500,
            content={"error": "Eroare la actualizarea datelor instituției"}
        )
    
    # Invalidează cache-ul
    if chat_id in _config_cache:
        del _config_cache[chat_id]
    
    # Reîncarcă config-ul
    config = get_cached_config(chat_id)
    
    return JSONResponse(content={
        "success": True,
        "message": "Datele instituției au fost actualizate",
        "config": config
    })

@app.put("/admin/tenant/{chat_id}/config")
async def update_tenant_config(chat_id: str, config_updates: dict):
    """Actualizează configurația unui tenant (prompt, model, setări UI, etc.)"""
    # Convertește chat_id la int
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        # Caută după name
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Actualizează în baza de date
    success = update_client_chat(
        chat_id=client_chat_id,
        name=config_updates.get("name"),
        model=config_updates.get("model"),
        prompt=config_updates.get("prompt"),
        chat_title=config_updates.get("chat_title"),
        chat_subtitle=config_updates.get("chat_subtitle"),
        chat_color=config_updates.get("chat_color"),
        is_active=config_updates.get("is_active")
    )
    
    if not success:
        return JSONResponse(
            status_code=500,
            content={"error": "Eroare la actualizarea configurației"}
        )
    
    # Invalidează cache-ul
    if chat_id in _config_cache:
        del _config_cache[chat_id]
    
    # Reîncarcă config-ul
    config = get_cached_config(chat_id)
    
    return JSONResponse(content={
        "success": True,
        "message": "Configurația a fost actualizată",
        "config": config
    })

@app.post("/admin/tenant/{chat_id}/rag/upload")
async def upload_rag_file(chat_id: str, file: UploadFile = File(...)):
    """Încarcă un fișier RAG pentru un tenant"""
    print(f"📤 Upload RAG pentru tenant {chat_id}, fișier: {file.filename if file.filename else 'N/A'}")
    
    config = get_cached_config(chat_id)
    if not config:
        print(f"❌ Config nu există pentru {chat_id}")
        return JSONResponse(
            status_code=404,
            content={"error": f"Chat configuration not found: {chat_id}"}
        )
    
    if not file.filename:
        print(f"❌ Fișier fără nume pentru {chat_id}")
        return JSONResponse(
            status_code=400,
            content={"error": "Fișierul trebuie să aibă un nume"}
        )
    
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    rag_dir = f"rag/{chat_id}"
    os.makedirs(rag_dir, exist_ok=True)
    
    # Salvează fișierul
    file_path = os.path.join(rag_dir, file.filename)
    try:
        content = await file.read()
        print(f"✅ Fișier citit: {len(content)} bytes")
        with open(file_path, "wb") as f:
            f.write(content)
        print(f"✅ Fișier salvat la: {file_path}")
    except Exception as e:
        print(f"❌ Eroare la salvarea fișierului: {e}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la salvarea fișierului: {str(e)}"}
        )
    
    # Extrage text (similar cu logica din reprocess_rag)
    text_content = ""
    try:
        print(f"📄 Încep extragerea textului din {file.filename}...")
        if file.filename.endswith('.pdf') and PDF_AVAILABLE:
            print(f"📄 Procesare PDF: {file.filename}")
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                print(f"📄 PDF are {len(pdf_reader.pages)} pagini")
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n--- Pagina {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        print(f"⚠️ Eroare la extragerea paginii {page_num + 1} din {file.filename}: {e}")
            if not text_content.strip():
                print(f"⚠️ PDF {file.filename} nu conține text extractibil (poate fi scanat)")
        elif file.filename.endswith(('.txt', '.md')):
            print(f"📄 Procesare text: {file.filename}")
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()
            except UnicodeDecodeError:
                print(f"⚠️ Eroare encoding UTF-8, încerc latin-1...")
                with open(file_path, "r", encoding="latin-1") as f:
                    text_content = f.read()
        elif file.filename.endswith(('.doc', '.docx')):
            print(f"📄 Procesare DOC/DOCX: {file.filename}")
            try:
                from docx import Document
                doc = Document(file_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content += para.text + "\n"
                # Extrage și din tabele
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_content += row_text + "\n"
            except ImportError:
                print(f"⚠️ python-docx nu este instalat. Pentru DOC/DOCX, rulează: pip install python-docx")
                text_content = f"[Fișier DOC/DOCX - instalează python-docx pentru extragere]"
        else:
            print(f"⚠️ Tip de fișier necunoscut: {file.filename}")
        
        print(f"✅ Text extras: {len(text_content)} caractere")
    except Exception as e:
        print(f"❌ Eroare la extragerea textului din {file.filename}: {e}")
        import traceback
        traceback.print_exc()
    
    # Convertește chat_id la int pentru DB
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Adaugă fișierul în baza de date cu conținutul text
    if text_content and text_content.strip():
        # Salvează sau actualizează fișierul în DB cu conținutul
        add_rag_file(client_chat_id, file.filename, text_content.strip())
        print(f"✅ Fișier RAG salvat în DB cu conținut: {file.filename} ({len(text_content)} caractere)")
    else:
        # Dacă nu s-a putut extrage text, salvează doar numele fișierului
        add_rag_file(client_chat_id, file.filename, None)
        print(f"⚠️ Nu s-a putut extrage text din {file.filename} (poate fi gol, scanat sau protejat) - salvat doar numele fișierului")
    
    # Actualizează vector store
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        if text_content and text_content.strip():
            rag_store.add_document(file.filename, text_content.strip())
            print(f"✅ Fișier RAG adăugat în vector store pentru tenant {tenant_id}")
        else:
            print(f"⚠️ Nu s-a adăugat în vector store (fără conținut text)")
    except Exception as e:
        print(f"⚠️ Eroare la actualizarea vector store pentru tenant {tenant_id}: {e}")
        import traceback
        traceback.print_exc()
    
    # Invalidează cache-ul
    if chat_id in _config_cache:
        del _config_cache[chat_id]
        if chat_id in _config_cache_timestamps:
            del _config_cache_timestamps[chat_id]
    
    return JSONResponse(content={
        "success": True,
        "message": f"Fișier {file.filename} încărcat cu succes" + ("" if text_content and text_content.strip() else " (fără conținut text extractibil)"),
        "filename": file.filename,
        "has_content": bool(text_content and text_content.strip()),
        "content_length": len(text_content) if text_content else 0
    })

@app.delete("/admin/tenant/{chat_id}/rag/{filename}")
async def delete_rag_file(chat_id: str, filename: str):
    """Șterge un fișier RAG pentru un tenant"""
    # Decode filename dacă este URL-encoded
    from urllib.parse import unquote
    filename = unquote(filename)
    
    print(f"🗑️ Ștergere RAG pentru tenant {chat_id}, fișier: {filename}")
    
    config = get_cached_config(chat_id)
    if not config:
        print(f"❌ Config nu există pentru {chat_id}")
        return JSONResponse(
            status_code=404,
            content={"error": f"Chat configuration not found: {chat_id}"}
        )
    
    tenant_id = get_tenant_id_from_chat_id(chat_id)
    
    # Șterge fișierul din folder
    file_path = f"rag/{chat_id}/{filename}"
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            print(f"✅ Fișier șters din folder: {file_path}")
        except Exception as e:
            print(f"⚠️ Eroare la ștergerea fișierului din folder: {e}")
    else:
        print(f"⚠️ Fișier nu există în folder: {file_path}")
    
    # Convertește chat_id la int pentru DB
    try:
        client_chat_id = int(chat_id)
    except ValueError:
        db_config = get_client_chat(chat_id)
        if not db_config:
            return JSONResponse(
                status_code=404,
                content={"error": f"Chat configuration not found: {chat_id}"}
            )
        client_chat_id = db_config.get("id")
    
    # Șterge din baza de date
    deleted = delete_rag_file(client_chat_id, filename)
    if deleted:
        print(f"✅ Fișier șters din DB: {filename}")
    else:
        print(f"⚠️ Fișier nu era în DB: {filename}")
    
    # Actualizează vector store
    try:
        rag_store = get_tenant_rag_store(tenant_id)
        rag_store.remove_document(filename)
        print(f"✅ Fișier RAG șters din vector store pentru tenant {tenant_id}")
    except Exception as e:
        print(f"⚠️ Eroare la actualizarea vector store pentru tenant {tenant_id}: {e}")
        import traceback
        traceback.print_exc()
    
    # Invalidează cache-ul
    if chat_id in _config_cache:
        del _config_cache[chat_id]
        if chat_id in _config_cache_timestamps:
            del _config_cache_timestamps[chat_id]
    
    return JSONResponse(content={
        "success": True,
        "message": f"Fișier {filename} șters cu succes"
    })

@app.get("/admin/tenants")
async def list_all_tenants():
    """Listează toți tenant-ii (pentru panoul de administrare)"""
    # Încarcă din baza de date
    db_tenants = list_all_client_chats()
    
    tenants = []
    for db_tenant in db_tenants:
        tenants.append({
            "id": str(db_tenant.get("id", "")),
            "tenant_id": str(db_tenant.get("id", "")),
            "name": db_tenant.get("name", "Unknown"),
            "model": db_tenant.get("model", "unknown"),
            "is_active": bool(db_tenant.get("is_active", True)),
            "created_at": db_tenant.get("created_at"),
            "updated_at": db_tenant.get("updated_at"),
            "institution": db_tenant.get("institution"),
            "rag_files_count": db_tenant.get("rag_files_count", 0),
            "chat_title": db_tenant.get("chat_title"),
            "chat_color": db_tenant.get("chat_color")
        })
    
    return JSONResponse(content={"tenants": tenants})

@app.post("/admin/tenant/create")
async def create_tenant(request: dict):
    """Creează un nou tenant/client chatbot"""
    try:
        name = request.get("name", "Chat nou")
        model = request.get("model", "gpt-oss:20b")
        prompt = request.get("prompt", "Ești asistentul Integra AI. Răspunde clar și politicos la întrebările utilizatorilor.")
        chat_title = request.get("chat_title", name)
        chat_subtitle = request.get("chat_subtitle", "Asistentul tău inteligent pentru găsirea informațiilor")
        chat_color = request.get("chat_color", "#3b82f6")
        
        # Creează chatbot-ul în baza de date
        client_chat_id = create_client_chat(
            name=name,
            model=model,
            prompt=prompt,
            chat_title=chat_title,
            chat_subtitle=chat_subtitle,
            chat_color=chat_color
        )
        
        if not client_chat_id:
            return JSONResponse(
                status_code=500,
                content={"error": "Eroare la crearea chatbot-ului în baza de date"}
            )
        
        # Reîncarcă config-ul din DB
        config = get_cached_config(str(client_chat_id))
        
        return JSONResponse(content={
            "success": True,
            "message": "Client creat cu succes",
            "tenant": {
                "id": str(client_chat_id),
                "tenant_id": str(client_chat_id),
                "name": name,
                "model": model,
                "is_active": True,
                "institution": None,
                "rag_files_count": 0,
                "chat_title": chat_title,
                "chat_color": chat_color
            }
        })
    except Exception as e:
        print(f"❌ Eroare la crearea tenant-ului: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": f"Eroare la crearea clientului: {str(e)}"}
        )

# ==================== ENDPOINT-URI PENTRU AUTENTIFICARE ====================

@app.post("/auth/register", response_model=TokenResponse)
async def register(request: RegisterRequest):
    """Înregistrare utilizator nou"""
    try:
        # Verifică dacă email-ul există deja
        existing_user = get_user(email=request.email)
        if existing_user:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Email-ul este deja înregistrat"
            )
        
        # Hash-uiește parola
        hashed_password = hash_password(request.password)
        
        # Creează utilizatorul
        user_id = create_user(
            name=request.name,
            email=request.email,
            password=hashed_password,
            role='user'
        )
        
        if not user_id:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Eroare la crearea contului"
            )
        
        # Obține utilizatorul creat
        user = get_user(user_id=user_id)
        if not user:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Eroare la obținerea datelor utilizatorului"
            )
        
        # Creează token JWT
        access_token = create_access_token(data={"sub": str(user_id), "email": request.email})
        
        # Elimină parola din răspuns
        user_response = {k: v for k, v in user.items() if k != 'password'}
        
        return TokenResponse(
            access_token=access_token,
            token_type="bearer",
            user=user_response
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Eroare la înregistrare: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Eroare la înregistrare: {str(e)}"
        )

@app.post("/auth/login", response_model=TokenResponse)
async def login(request: LoginRequest):
    """Autentificare utilizator"""
    try:
        # Obține utilizatorul după email
        user = get_user(email=request.email)
        if not user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Email sau parolă incorectă"
            )
        
        # Verifică parola
        if not verify_password(request.password, user['password']):
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Email sau parolă incorectă"
            )
        
        # Creează token JWT
        access_token = create_access_token(data={"sub": str(user['id']), "email": user['email']})
        
        # Elimină parola din răspuns
        user_response = {k: v for k, v in user.items() if k != 'password'}
        
        return TokenResponse(
            access_token=access_token,
            token_type="bearer",
            user=user_response
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Eroare la login: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Eroare la autentificare: {str(e)}"
        )

@app.get("/auth/me")
async def get_current_user_info(current_user: dict = Depends(get_current_user)):
    """Obține informațiile utilizatorului curent"""
    if current_user is None:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Nu ești autentificat"
        )
    return JSONResponse(content=current_user)

@app.post("/auth/logout")
async def logout():
    """Logout (în realitate, clientul trebuie să șteargă token-ul)"""
    return JSONResponse(content={"message": "Logout reușit"})

